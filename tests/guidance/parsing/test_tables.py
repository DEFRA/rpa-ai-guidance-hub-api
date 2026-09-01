"""Turning a Word table into the Markdown it says.

A cell is a little document of its own, so most of these cases are about what
survives being read as one: a link, a mark, a second paragraph, a list. The rest are
about the grid, where Markdown can express less than Word can - it spans neither
columns nor rows, and a pipe row cannot hold a line break at all.

All fixture text is invented, as everywhere in this package.
"""

from collections.abc import Callable
from typing import Any

from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

from app.guidance.parsing import parser

HEADING = "Applying"


def _table_content(
    docx_bytes: Callable[..., bytes], build: Callable[[Any], None]
) -> str:
    """The Markdown of the one section `build` writes its table into.

    Named for this file rather than `_content`, which two of the sibling test
    modules also define with a signature of their own. Nothing at runtime confuses
    them - pytest imports each file as its own module - but the package carries no
    __init__.py, so an editor indexing the directory has no boundary to tell it that
    three same-named helpers are three different functions.
    """

    def with_heading(document: Any) -> None:
        document.add_heading(HEADING, level=1)
        build(document)

    return parser.parse_docx(docx_bytes(with_heading)).sections[0].content


def _cell(table, row: int, column: int):
    """One cell as Word wrote it, without python-docx's grid expansion."""
    return table._tbl.findall(qn("w:tr"))[row].findall(qn("w:tc"))[column]


def _cell_property(cell, name: str):
    """Add one cell property, creating the w:tcPr it has to live in."""
    properties = cell.find(qn("w:tcPr"))
    if properties is None:
        properties = OxmlElement("w:tcPr")
        # w:tcPr is the first child of a w:tc wherever it appears.
        cell.insert(0, properties)

    element = OxmlElement(name)
    properties.append(element)
    return element


def _spans(cell, columns: int) -> None:
    """Make a cell cover several grid columns, as a merged heading does."""
    _cell_property(cell, "w:gridSpan").set(qn("w:val"), str(columns))


def _merges(cell, *, origin: bool) -> None:
    """Mark a cell as starting a vertical merge, or as carrying one on."""
    element = _cell_property(cell, "w:vMerge")
    if origin:
        element.set(qn("w:val"), "restart")


def _remove(cell) -> None:
    """Take a cell out, as Word does when its neighbour spans its column."""
    cell.getparent().remove(cell)


def _hyperlink(paragraph, text: str, url: str) -> None:
    """Append a w:hyperlink pointing outside the document."""
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True))
    link.append(paragraph.add_run(text)._r)
    paragraph._p.append(link)


class TestTheGrid:
    def test_a_table_renders_as_a_pipe_table_headed_by_its_first_row(self, docx_bytes):
        """GFM needs a header row and no real guide declares one, so the first row
        is it - which is what every one of them makes bold."""

        def build(document):
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Case name"
            table.cell(0, 1).text = "Action to take"
            table.cell(1, 0).text = "Withdrawal"
            table.cell(1, 1).text = "Close the case"

        assert _table_content(docx_bytes, build) == (
            "| Case name | Action to take |\n"
            "| --- | --- |\n"
            "| Withdrawal | Close the case |"
        )

    def test_whitespace_inside_a_cell_is_collapsed(self, docx_bytes):
        """The editor collapses it on every cell it writes, so a cell left as Word
        spelled it would come back changed by the first save."""

        def build(document):
            table = document.add_table(rows=2, cols=1)
            table.cell(0, 0).text = "Case  name"
            table.cell(1, 0).text = "  Withdrawal  "

        assert _table_content(docx_bytes, build) == (
            "| Case name |\n| --- |\n| Withdrawal |"
        )

    def test_a_spanned_cell_fills_one_column_and_leaves_the_rest_empty(
        self, docx_bytes
    ):
        """Markdown has no column span. Repeating the label instead would put words
        in the output that the document does not say."""

        def build(document):
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Historic"
            table.cell(1, 0).text = "Parcel"
            table.cell(1, 1).text = "Area"
            _remove(_cell(table, 0, 1))
            _spans(_cell(table, 0, 0), 2)

        assert _table_content(docx_bytes, build) == (
            "| Historic |  |\n| --- | --- |\n| Parcel | Area |"
        )

    def test_a_merged_cell_says_nothing_where_it_carries_on(self, docx_bytes):
        """A continuation is not on the page: Word draws the cell above through it.
        Reading it as a cell of its own is what repeated 17 of them in one guide."""

        def build(document):
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Withdrawal"
            table.cell(0, 1).text = "Close the case"
            table.cell(1, 0).text = "Withdrawal"
            table.cell(1, 1).text = "Tell the agent"
            _merges(_cell(table, 0, 0), origin=True)
            _merges(_cell(table, 1, 0), origin=False)

        assert _table_content(docx_bytes, build).endswith("|  | Tell the agent |")

    def test_an_empty_cell_keeps_its_column(self, docx_bytes):
        def build(document):
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Case name"
            table.cell(0, 1).text = "Action to take"
            table.cell(1, 0).text = "Withdrawal"

        assert _table_content(docx_bytes, build).endswith("| Withdrawal |  |")

    def test_a_table_with_no_rows_says_nothing(self, docx_bytes):
        def build(document):
            document.add_table(rows=0, cols=2)
            document.add_paragraph("Send the form")

        assert _table_content(docx_bytes, build) == "Send the form"


class TestWhatACellHolds:
    def test_a_link_in_a_cell_survives(self, docx_bytes):
        """32 URLs across the two real guides live in cells, and reading a cell as
        text is what lost every one of them."""

        def build(document):
            table = document.add_table(rows=2, cols=1)
            table.cell(0, 0).text = "Guidance"
            _hyperlink(
                table.cell(1, 0).paragraphs[0],
                "The claims guide",
                "https://example.org",
            )

        assert _table_content(docx_bytes, build).endswith(
            "| [The claims guide](https://example.org) |"
        )

    def test_marks_in_a_cell_survive(self, docx_bytes):
        def build(document):
            table = document.add_table(rows=2, cols=1)
            table.cell(0, 0).text = "Case name"
            table.cell(1, 0).paragraphs[0].add_run("Withdrawal").bold = True

        assert _table_content(docx_bytes, build).endswith("| **Withdrawal** |")

    def test_a_cell_of_two_paragraphs_stays_on_one_row(self, docx_bytes):
        """A newline anywhere in a row ends the row and orphans what follows it,
        which is how 10 rows were lost in one guide."""

        def build(document):
            table = document.add_table(rows=2, cols=1)
            table.cell(0, 0).text = "Action to take"
            cell = table.cell(1, 0)
            cell.text = "Close the case"
            cell.add_paragraph("Then tell the agent")

        assert _table_content(docx_bytes, build).endswith(
            "| Close the case<br>Then tell the agent |"
        )

    def test_a_list_in_a_cell_keeps_its_markers(self, docx_bytes):
        """A Markdown list cannot live in a pipe row, so the markers are all that is
        left to say the items were a list."""

        def build(document):
            table = document.add_table(rows=2, cols=1)
            table.cell(0, 0).text = "Action to take"
            cell = table.cell(1, 0)
            cell.text = "Do both:"
            cell.add_paragraph("Close the case", style="List Bullet")
            cell.add_paragraph("Tell the agent", style="List Bullet")

        assert _table_content(docx_bytes, build).endswith(
            "| Do both:<br>- Close the case<br>- Tell the agent |"
        )

    def test_a_line_break_in_a_cell_does_not_end_the_row(self, docx_bytes):
        """The break's own backslash goes with it: it marks a break there is no
        longer room for."""

        def build(document):
            table = document.add_table(rows=2, cols=1)
            table.cell(0, 0).text = "Action to take"
            paragraph = table.cell(1, 0).paragraphs[0]
            paragraph.add_run("Close the case")
            paragraph.add_run().add_break()
            paragraph.add_run("Tell the agent")

        assert _table_content(docx_bytes, build).endswith(
            "| Close the case<br>Tell the agent |"
        )

    def test_a_pipe_in_a_cell_is_escaped(self, docx_bytes):
        """A pipe ends a cell wherever it appears, so one meant as text has to say
        so or it silently adds a column."""

        def build(document):
            table = document.add_table(rows=2, cols=1)
            table.cell(0, 0).text = "Case name"
            table.cell(1, 0).text = "Withdrawal | reject"

        assert _table_content(docx_bytes, build).endswith("| Withdrawal \\| reject |")

    def test_the_pipe_rule_and_the_text_rule_compose(self, docx_bytes):
        """The pipe is a property of the row and the rest a property of the text, so
        each is escaped where it is known about and neither undoes the other."""

        def build(document):
            table = document.add_table(rows=2, cols=1)
            table.cell(0, 0).text = "Case name"
            table.cell(1, 0).text = "<Claim ID> | [draft]"

        assert _table_content(docx_bytes, build).endswith(
            "| &lt;Claim ID&gt; \\| \\[draft\\] |"
        )


class TestCallouts:
    def test_a_one_cell_table_is_a_blockquote(self, docx_bytes):
        """Word has no callout box, so a table of one cell is what it uses instead."""

        def build(document):
            document.add_table(rows=1, cols=1).cell(0, 0).text = "Check the version"

        assert _table_content(docx_bytes, build) == "> Check the version"

    def test_a_callout_keeps_its_paragraphs_apart(self, docx_bytes):
        """All nine callouts in the two real guides have more than one paragraph,
        and reading the cell as one string put them on one line."""

        def build(document):
            cell = document.add_table(rows=1, cols=1).cell(0, 0)
            cell.text = "Agreement holder:"
            cell.add_paragraph("Permission level:")

        assert _table_content(docx_bytes, build) == (
            "> Agreement holder:\n>\n> Permission level:"
        )

    def test_an_empty_callout_says_nothing(self, docx_bytes):
        def build(document):
            document.add_table(rows=1, cols=1)
            document.add_paragraph("Send the form")

        assert _table_content(docx_bytes, build) == "Send the form"


class TestWhereATableGoes:
    def test_a_table_closes_an_open_list_run(self, docx_bytes):
        """Two runs either side of a table are two lists, not one."""

        def build(document):
            document.add_paragraph("Send the form", style="List Bullet")
            document.add_table(rows=1, cols=1).cell(0, 0).text = "Check the version"
            document.add_paragraph("Keep a copy", style="List Bullet")

        assert _table_content(docx_bytes, build) == (
            "- Send the form\n\n> Check the version\n\n- Keep a copy"
        )

    def test_a_table_files_under_the_section_open_at_the_time(self, docx_bytes):
        def build(document):
            document.add_heading(HEADING, level=1)
            document.add_table(rows=1, cols=1).cell(0, 0).text = "Check the version"
            document.add_heading("Assessing", level=1)
            document.add_table(rows=1, cols=1).cell(0, 0).text = "Check the claim"

        sections = parser.parse_docx(docx_bytes(build)).sections
        assert [section.content for section in sections] == [
            "> Check the version",
            "> Check the claim",
        ]

    def test_a_table_ahead_of_every_heading_is_dropped(self, docx_bytes):
        """It has no section to belong to, exactly as a paragraph there has none."""

        def build(document):
            document.add_table(rows=1, cols=1).cell(0, 0).text = "Check the version"

        assert parser.parse_docx(docx_bytes(build)).sections == []
