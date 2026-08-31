"""Turning heading paragraphs into numbered sections.

Word does not put the number in the text. Both real guides attach the numbering to
the heading *styles*, so Word generates "4.3.1.1" as it renders the page and the
paragraph itself says only "Split". Every number here is therefore derived from the
heading structure alone, and these cases pin down how.

Levels are read as relative, so most of the interesting cases are documents whose
headings do not descend one tidy level at a time.
"""

import pytest
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

from app.guidance.parsing import parser


def _bookmark_start(name: str):
    """A w:bookmarkStart element, as Word opens a bookmark."""
    element = OxmlElement("w:bookmarkStart")
    element.set(qn("w:id"), "1")
    element.set(qn("w:name"), name)
    return element


def _bookmark_in(paragraph, *names: str) -> None:
    """Open bookmarks inside a paragraph, where Word usually writes them."""
    for name in names:
        paragraph._p.append(_bookmark_start(name))


def _bookmark_before(paragraph, name: str) -> None:
    """Open a bookmark at body level, as a sibling just ahead of a paragraph."""
    paragraph._p.addprevious(_bookmark_start(name))


def _anchor_link(paragraph, text: str, anchor: str) -> None:
    """Append a link to a bookmark inside the document."""
    link = OxmlElement("w:hyperlink")
    link.set(qn("w:anchor"), anchor)
    link.append(paragraph.add_run(text)._r)
    paragraph._p.append(link)


def _outline(source: bytes) -> list[tuple[str, str]]:
    """The parsed sections as (number, heading) pairs, in document order."""
    return [
        (section.number, section.heading)
        for section in parser.parse_docx(source).sections
    ]


class TestBuildingTheSectionTree:
    def test_headings_nest_by_level(self, docx_bytes):
        def build(document):
            document.add_heading("Applying", level=1)
            document.add_heading("Eligibility", level=2)
            document.add_heading("Land parcels", level=3)

        assert _outline(docx_bytes(build)) == [
            ("1", "Applying"),
            ("1.1", "Eligibility"),
            ("1.1.1", "Land parcels"),
        ]

    def test_each_parent_counts_its_own_children(self, docx_bytes):
        """The second parent's children restart at 1 rather than carrying on."""

        def build(document):
            document.add_heading("Applying", level=1)
            document.add_heading("Eligibility", level=2)
            document.add_heading("Evidence", level=2)
            document.add_heading("Assessing", level=1)
            document.add_heading("Checks", level=2)

        assert _outline(docx_bytes(build)) == [
            ("1", "Applying"),
            ("1.1", "Eligibility"),
            ("1.2", "Evidence"),
            ("2", "Assessing"),
            ("2.1", "Checks"),
        ]

    def test_returning_to_a_shallower_level_resumes_the_outer_count(self, docx_bytes):
        """Closing two levels at once must not lose the top-level count."""

        def build(document):
            document.add_heading("Applying", level=1)
            document.add_heading("Eligibility", level=2)
            document.add_heading("Land parcels", level=3)
            document.add_heading("Assessing", level=1)

        assert _outline(docx_bytes(build))[-1] == ("2", "Assessing")

    def test_a_document_opening_at_heading_2_still_starts_at_one(self, docx_bytes):
        """Level is relative to the document, not to Word's numbering."""

        def build(document):
            document.add_heading("Applying", level=2)
            document.add_heading("Assessing", level=2)

        assert _outline(docx_bytes(build)) == [
            ("1", "Applying"),
            ("2", "Assessing"),
        ]

    def test_a_skipped_level_nests_one_deep(self, docx_bytes):
        """A jump from 1 to 3 is a child, not a gap: there is no "1.0.1"."""

        def build(document):
            document.add_heading("Applying", level=1)
            document.add_heading("Land parcels", level=3)
            document.add_heading("Evidence", level=3)

        assert _outline(docx_bytes(build)) == [
            ("1", "Applying"),
            ("1.1", "Land parcels"),
            ("1.2", "Evidence"),
        ]


class TestWhatDoesNotOpenASection:
    def test_text_before_the_first_heading_opens_no_section(self, docx_bytes):
        def build(document):
            document.add_paragraph("This guide covers the whole claim.")
            document.add_heading("Applying", level=1)

        assert _outline(docx_bytes(build)) == [("1", "Applying")]

    def test_a_heading_with_no_text_is_not_a_section(self, docx_bytes):
        """An empty heading is spacing. Numbering it invents a section."""

        def build(document):
            document.add_heading("Applying", level=1)
            document.add_heading("   ", level=1)
            document.add_heading("Assessing", level=1)

        assert _outline(docx_bytes(build)) == [
            ("1", "Applying"),
            ("2", "Assessing"),
        ]

    @pytest.mark.parametrize(
        "style_name",
        [
            pytest.param("Heading Box", id="no-level"),
            pytest.param("Heading 2 Box", id="level-then-more-name"),
            pytest.param("Box Heading 2", id="more-name-then-level"),
            pytest.param("Heading 0", id="level-zero"),
        ],
    )
    def test_a_style_only_named_like_a_heading_is_body_text(
        self, docx_bytes, in_style, style_name
    ):
        """A heading style is named for the word and a level from 1, and nothing else.

        Templates carry plenty of styles built around the word - a heading for a
        table, a heading for an annex - and none of them is a body heading. Reading
        a level out of one anyway would also let a "Heading 0" outrank the document
        itself, leaving a section with nothing to hang from.
        """

        def build(document):
            document.add_heading("Applying", level=1)
            in_style(document, "Deadline reminder", style_name)

        assert _outline(docx_bytes(build)) == [("1", "Applying")]

    def test_a_document_with_no_headings_has_no_sections(self, docx_bytes):
        def build(document):
            document.add_paragraph("This guide covers the whole claim.")

        assert _outline(docx_bytes(build)) == []


class TestWhereContentGoes:
    def test_a_paragraph_belongs_to_the_section_opened_most_recently(self, docx_bytes):
        """Whatever its depth: prose under 1.1 is 1.1's, not section 1's."""

        def build(document):
            document.add_heading("Applying", level=1)
            document.add_paragraph("Read this section first.")
            document.add_heading("Eligibility", level=2)
            document.add_paragraph("Apply before the deadline.")

        sections = parser.parse_docx(docx_bytes(build)).sections

        assert [section.content for section in sections] == [
            "Read this section first.",
            "Apply before the deadline.",
        ]

    def test_nothing_ahead_of_the_first_heading_becomes_content(self, docx_bytes):
        """In both real guides everything there is the cover page and the contents.

        25 and 30 blocks of it, and not one line of body prose - so it is left out
        rather than filed under a section it does not belong to.
        """

        def build(document):
            document.add_paragraph("Example Grant Scheme Guide", style="Title")
            document.add_paragraph("Printed on recycled paper.")
            document.add_heading("Applying", level=1)
            document.add_paragraph("Apply before the deadline.")

        document = parser.parse_docx(docx_bytes(build))

        assert document.sections[0].content == "Apply before the deadline."
        assert "recycled" not in document.markdown()

    def test_a_contents_entry_after_a_heading_is_still_not_content(
        self, docx_bytes, in_style
    ):
        """Word regenerates a contents page, so its entries are never prose.

        They sit ahead of every heading in both real guides and so never reach the
        content walk at all, but a contents page opening with a heading of its own
        would otherwise file all 23 entries as that section's text.
        """

        def build(document):
            document.add_heading("Contents", level=1)
            in_style(document, "1 Applying ....... 3", "TOC 1")
            document.add_heading("Applying", level=1)
            document.add_paragraph("Apply before the deadline.")

        sections = parser.parse_docx(docx_bytes(build)).sections

        assert [section.content for section in sections] == [
            "",
            "Apply before the deadline.",
        ]


class TestRenderingTheSections:
    def test_a_section_is_headed_by_its_number_at_its_own_depth(self, docx_bytes):
        """The hash count follows the parent chain, so a skipped level stays sane."""

        def build(document):
            document.add_heading("Applying", level=1)
            document.add_heading("Land parcels", level=3)

        markdown = parser.parse_docx(docx_bytes(build)).markdown()

        assert "## 1 Applying" in markdown
        assert "### 1.1 Land parcels" in markdown


class TestAppendices:
    """An annex is a section Word gives no number and no outline level.

    Both are carried by a paragraph style instead - CS names its own `Appendix`,
    pulled into the contents by `TOC \\t "Appendix,1"` - so the style name is the
    only thing that says an annex is starting. Left unrecognised, an annex is not a
    section at all and its text is filed under whichever one was last open.
    """

    def test_an_annex_style_opens_a_section_and_takes_its_own_content(
        self, docx_bytes, in_style
    ):
        """The defect this feature exists to fix, stated directly."""

        def build(document):
            document.add_heading("Applying", level=1)
            document.add_paragraph("Apply before the deadline.")
            in_style(document, "Annex A - Case types", "Appendix")
            document.add_paragraph("The table lists the cases you may see.")

        sections = parser.parse_docx(docx_bytes(build)).sections

        assert [(s.number, s.heading) for s in sections] == [
            ("1", "Applying"),
            ("A", "Annex A - Case types"),
        ]
        assert sections[0].content == "Apply before the deadline."
        assert sections[1].content == "The table lists the cases you may see."

    def test_the_letter_comes_from_position_not_from_the_heading(
        self, docx_bytes, in_style
    ):
        """The author's own designation is text, and can say anything at all."""

        def build(document):
            in_style(document, "Annex Z - Case types", "Appendix")
            in_style(document, "Annex Q - Reject a claim", "Appendix")

        assert _outline(docx_bytes(build)) == [
            ("A", "Annex Z - Case types"),
            ("B", "Annex Q - Reject a claim"),
        ]

    def test_annexes_and_numbered_sections_are_counted_apart(
        self, docx_bytes, in_style
    ):
        """Neither count disturbs the other: an annex is not section 2."""

        def build(document):
            document.add_heading("Applying", level=1)
            in_style(document, "Annex A - Case types", "Appendix")
            document.add_heading("Assessing", level=1)

        assert _outline(docx_bytes(build)) == [
            ("1", "Applying"),
            ("A", "Annex A - Case types"),
            ("2", "Assessing"),
        ]

    def test_a_heading_under_an_annex_nests_beneath_its_letter(
        self, docx_bytes, in_style
    ):
        """An annex is a top level section, so a deeper heading nests beneath it.

        A Heading 1 after an annex is a sibling instead, and is numbered - which is
        the case above, and the reason the two counts are kept apart.
        """

        def build(document):
            in_style(document, "Annex A - Case types", "Appendix")
            document.add_heading("Rejected claims", level=2)
            document.add_heading("Evidence", level=3)

        assert _outline(docx_bytes(build)) == [
            ("A", "Annex A - Case types"),
            ("A.1", "Rejected claims"),
            ("A.1.1", "Evidence"),
        ]

    @pytest.mark.parametrize(
        "style_name",
        [
            pytest.param("Appendix", id="appendix"),
            pytest.param("Annex Heading", id="annex"),
            pytest.param("Schedule Title", id="schedule"),
        ],
    )
    def test_every_name_a_template_gives_the_style_is_recognised(
        self, docx_bytes, in_style, style_name
    ):
        """One template's word for it is another's; all three mean an annex."""

        def build(document):
            in_style(document, "Case types", style_name)

        assert _outline(docx_bytes(build)) == [("A", "Case types")]

    def test_an_annex_with_no_text_is_not_a_section(self, docx_bytes, in_style):
        """Spacing in an annex style is still spacing, as it is for a heading."""

        def build(document):
            in_style(document, "   ", "Appendix")
            in_style(document, "Annex A - Case types", "Appendix")

        assert _outline(docx_bytes(build)) == [("A", "Annex A - Case types")]


class TestBookmarksACrossReferenceCanNameASectionBy:
    """Word writes a cross-reference as a bookmark name, which no renderer knows.

    The map from name to section exists only while the .docx is open, so if the
    parser does not build it nothing downstream can. A name is claimed only where
    it marks the start of a section, that being the whole of what has a number.
    """

    def test_a_bookmark_on_a_heading_names_that_section(self, docx_bytes):
        def build(document):
            document.add_heading("Applying", level=1)
            _bookmark_in(document.paragraphs[-1], "_Applying")

        bookmarks = parser.parse_docx(docx_bytes(build)).bookmarks

        assert bookmarks["_Applying"].number == "1"

    def test_a_bookmark_opened_ahead_of_a_heading_names_the_section_it_opens(
        self, docx_bytes
    ):
        """Word writes a bookmark wrapping a heading as a sibling before it.

        Attributing it to the paragraph already walked would name the section it
        sits after, which is the one the reader is being sent away from.
        """

        def build(document):
            document.add_heading("Applying", level=1)
            document.add_heading("Eligibility", level=2)
            _bookmark_before(document.paragraphs[-1], "_Eligibility")

        bookmarks = parser.parse_docx(docx_bytes(build)).bookmarks

        assert bookmarks["_Eligibility"].number == "1.1"

    def test_every_name_on_one_heading_names_it(self, docx_bytes):
        """Word keeps its own _Toc bookmark beside the author's on ~20 headings."""

        def build(document):
            document.add_heading("Applying", level=1)
            _bookmark_in(document.paragraphs[-1], "_Toc178312", "_Applying")

        bookmarks = parser.parse_docx(docx_bytes(build)).bookmarks

        assert bookmarks["_Toc178312"] is bookmarks["_Applying"]

    def test_a_bookmark_marking_no_section_start_is_not_claimed(self, docx_bytes):
        """A bookmark on a list bullet has no number it could be resolved to.

        Leaving it unclaimed keeps the raw name Word wrote, which is a link that
        goes nowhere rather than one that confidently goes to the wrong section.
        """

        def build(document):
            document.add_heading("Applying", level=1)
            document.add_paragraph("Send the form before the deadline.")
            _bookmark_in(document.paragraphs[-1], "_Deadline")

        bookmarks = parser.parse_docx(docx_bytes(build)).bookmarks

        assert "_Deadline" not in bookmarks

    def test_a_name_the_next_paragraph_does_not_claim_is_dropped(self, docx_bytes):
        """A bookmark ahead of ordinary prose marks that prose, not a later heading.

        Carrying it on would send the cross-reference to whichever section happened
        to start next, which is the wrong-section defect in another disguise.
        """

        def build(document):
            document.add_heading("Applying", level=1)
            _bookmark_before(document.add_paragraph("Send the form."), "_Deadline")
            document.add_heading("Payment", level=1)

        bookmarks = parser.parse_docx(docx_bytes(build)).bookmarks

        assert "_Deadline" not in bookmarks

    def test_a_cross_reference_renders_as_the_number_of_its_section(self, docx_bytes):
        """Forward references are why this is resolved at render time and not sooner.

        The section a link points at does not exist yet when the link is parsed, and
        its number is not final until the whole document has been walked.
        """

        def build(document):
            document.add_heading("Applying", level=1)
            _anchor_link(document.add_paragraph("Continue to "), "Payment", "_Payment")
            document.add_heading("Payment", level=1)
            _bookmark_in(document.paragraphs[-1], "_Payment")

        rendered = parser.parse_docx(docx_bytes(build)).markdown()

        assert "[Payment](#2)" in rendered

    def test_a_cross_reference_to_an_annex_resolves_to_its_letter(
        self, docx_bytes, in_style
    ):
        """The two annex links the PoC sent to the wrong section entirely.

        Their bookmarks sit on an annex, so they resolve only once an annex is a
        section - which is why they were left raw rather than guessed at.
        """

        def build(document):
            document.add_heading("Applying", level=1)
            _anchor_link(document.add_paragraph("See "), "Case types", "AnnexA")
            in_style(document, "Annex A - Case types", "Appendix")
            _bookmark_in(document.paragraphs[-1], "AnnexA")

        rendered = parser.parse_docx(docx_bytes(build)).markdown()

        assert "[Case types](#A)" in rendered
