"""Turning a paragraph's runs into the Markdown it says.

Word splits a paragraph into runs wherever anything at all changes, and puts the
space between two words inside whichever run it likes, so most of these cases are
about putting a paragraph back together rather than about any one mark.

Each case is parsed through a document with a heading, because content only reaches
the output as a section's content.
"""

from collections.abc import Callable

import pytest
from docx.document import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from app.guidance.parsing import parser

HEADING = "Applying"

# The `docx_bytes` fixture's contract, which its docstring otherwise states only in
# prose: a test hands it a callable that populates a document and gets the .docx
# bytes back. Naming it here is also what stops an editor inferring these helpers'
# parameters structurally and guessing wrong.
Build = Callable[[Document], None]
DocxBytes = Callable[[Build], bytes]


def _content(docx_bytes: DocxBytes, build: Build) -> str:
    """The Markdown of the one section `build` writes its paragraphs into."""

    def with_heading(document: Document) -> None:
        document.add_heading(HEADING, level=1)
        build(document)

    return parser.parse_docx(docx_bytes(with_heading)).sections[0].content


def _colour(run: Run, value: str) -> None:
    """Colour a run by raw w:val, including values python-docx will not write."""
    element = OxmlElement("w:color")
    element.set(qn("w:val"), value)
    run._r.get_or_add_rPr().append(element)


def _hyperlink(
    paragraph: Paragraph, *texts: str, url: str = "", anchor: str = ""
) -> list[Run]:
    """Append a w:hyperlink holding one run per text, as Word splits anchor text."""
    link = OxmlElement("w:hyperlink")
    if url:
        link.set(
            qn("r:id"), paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
        )
    if anchor:
        link.set(qn("w:anchor"), anchor)

    runs = [paragraph.add_run(text) for text in texts]
    for run in runs:
        link.append(run._r)
    paragraph._p.append(link)
    return runs


def _fld_char(paragraph: Paragraph, char_type: str) -> None:
    """Append a run holding one field boundary."""
    element = OxmlElement("w:fldChar")
    element.set(qn("w:fldCharType"), char_type)
    paragraph.add_run()._r.append(element)


def _instr_text(paragraph: Paragraph, text: str) -> None:
    """Append a run holding part of a field's instruction."""
    element = OxmlElement("w:instrText")
    element.text = text
    paragraph.add_run()._r.append(element)


def _field(
    paragraph: Paragraph,
    *instructions: str,
    result: str | None = None,
    closed: bool = True,
) -> None:
    """Append a Word field: begin, instruction, separate, result, end.

    One run per instruction string, as Word splits a long instruction. Passing no
    result leaves out the separate too, which is a field Word never rendered;
    `closed=False` leaves out the end, which is a paragraph Word stopped mid-field.
    """
    _fld_char(paragraph, "begin")
    for instruction in instructions:
        _instr_text(paragraph, instruction)
    if result is not None:
        _fld_char(paragraph, "separate")
        paragraph.add_run(result)
    if closed:
        _fld_char(paragraph, "end")


def _break(run: Run, break_type: str | None = None) -> None:
    """Add a w:br to a run, of the given type or of none at all."""
    element = OxmlElement("w:br")
    if break_type is not None:
        element.set(qn("w:type"), break_type)
    run._r.append(element)


class TestPuttingRunsBackTogether:
    def test_runs_wearing_the_same_marks_are_one_span(self, docx_bytes):
        """Word splits a bold phrase at a proofing boundary; it is still one phrase."""

        def build(document):
            paragraph = document.add_paragraph()
            paragraph.add_run("Is ").bold = True
            paragraph.add_run("your claim valid").bold = True

        assert _content(docx_bytes, build) == "**Is your claim valid**"

    def test_a_marker_never_wraps_the_space_word_kept_inside_it(self, docx_bytes):
        """The defect this whole module exists to avoid.

        Word puts the trailing space inside the bold run. `**Note: **read` is not
        emphasis under CommonMark - the asterisks are shown to the reader - and an
        editor that tidies the space outwards deletes it, welding the words together.
        """

        def build(document):
            paragraph = document.add_paragraph()
            paragraph.add_run("Note: ").bold = True
            paragraph.add_run("read the rules before applying.")

        assert (
            _content(docx_bytes, build) == "**Note:** read the rules before applying."
        )

    def test_a_span_of_nothing_but_space_carries_no_markers(self, docx_bytes):
        """Otherwise the space between two bold phrases renders as a literal ****."""

        def build(document):
            paragraph = document.add_paragraph()
            paragraph.add_run("Check")
            paragraph.add_run(" ").bold = True
            paragraph.add_run("again")

        assert _content(docx_bytes, build) == "Check again"

    def test_differently_marked_runs_stay_separate(self, docx_bytes):
        def build(document):
            paragraph = document.add_paragraph()
            paragraph.add_run("Check ").bold = True
            paragraph.add_run("carefully").italic = True

        assert _content(docx_bytes, build) == "**Check** *carefully*"

    def test_each_paragraph_is_its_own_block(self, docx_bytes):
        def build(document):
            document.add_paragraph("Apply before the deadline.")
            document.add_paragraph("Keep a copy of the form.")

        assert _content(docx_bytes, build) == (
            "Apply before the deadline.\n\nKeep a copy of the form."
        )

    def test_a_paragraph_saying_nothing_is_not_a_block(self, docx_bytes):
        """A blank paragraph is spacing, and would otherwise leave a stray gap."""

        def build(document):
            document.add_paragraph("Apply before the deadline.")
            document.add_paragraph("   ")
            document.add_paragraph("Keep a copy of the form.")

        assert _content(docx_bytes, build) == (
            "Apply before the deadline.\n\nKeep a copy of the form."
        )


class TestTheMarksARunCarries:
    @pytest.mark.parametrize(
        ("mark", "expected"),
        [
            pytest.param("bold", "**Overdue**", id="bold"),
            pytest.param("italic", "*Overdue*", id="italic"),
            pytest.param("underline", "<u>Overdue</u>", id="underline"),
            pytest.param("strike", "~~Overdue~~", id="strikethrough"),
            pytest.param("superscript", "<sup>Overdue</sup>", id="superscript"),
            pytest.param("subscript", "<sub>Overdue</sub>", id="subscript"),
        ],
    )
    def test_a_mark_becomes_markdown_or_the_html_markdown_allows(
        self, docx_bytes, mark, expected
    ):
        """Underline, superscript and subscript have no Markdown of their own."""

        def build(document):
            run = document.add_paragraph().add_run("Overdue")
            setattr(run.font, mark, True)

        assert _content(docx_bytes, build) == expected

    @pytest.mark.parametrize(
        "mark",
        [
            pytest.param("bold", id="toggle-set-to-zero"),
            pytest.param("underline", id="line-style-set-to-none"),
        ],
    )
    def test_a_mark_turned_off_is_not_a_mark(self, docx_bytes, mark):
        """Turning a mark off writes it out as off rather than removing it.

        The two are written differently - a toggle takes w:val="0", underline names
        the line "none" - so presence alone would read either as on.
        """

        def build(document):
            run = document.add_paragraph().add_run("Overdue")
            setattr(run.font, mark, False)

        assert _content(docx_bytes, build) == "Overdue"

    def test_marks_nest_in_a_fixed_order(self, docx_bytes):
        def build(document):
            run = document.add_paragraph().add_run("Overdue")
            run.font.bold = True
            run.font.italic = True
            run.font.underline = True

        assert _content(docx_bytes, build) == "***<u>Overdue</u>***"


class TestHyperlinks:
    def test_a_link_split_across_runs_is_one_link(self, docx_bytes):
        """w:hyperlink is the unit, however many runs Word broke its text into."""

        def build(document):
            _hyperlink(
                document.add_paragraph(),
                "Guid",
                "ance page",
                url="https://example.org/guidance",
            )

        assert _content(docx_bytes, build) == (
            "[Guidance page](https://example.org/guidance)"
        )

    def test_adjacent_links_to_the_same_target_are_one_link(self, docx_bytes):
        """Word writes a phrase broken by an edit as two w:hyperlink siblings.

        The target is a mark like any other, so the rule that merges adjacent runs
        wearing the same marks already makes these one link. Without it the reader
        is shown the same destination twice in a row, mid-phrase.
        """

        def build(document):
            paragraph = document.add_paragraph()
            _hyperlink(paragraph, "Basic ", url="https://example.org/guidance")
            _hyperlink(paragraph, "navigation", url="https://example.org/guidance")

        assert _content(docx_bytes, build) == (
            "[Basic navigation](https://example.org/guidance)"
        )

    def test_adjacent_links_to_different_targets_stay_apart(self, docx_bytes):
        """The guard on merging: a shared target is what makes two links one."""

        def build(document):
            paragraph = document.add_paragraph()
            _hyperlink(paragraph, "Claim form", url="https://example.org/claim")
            _hyperlink(paragraph, "Payment dates", url="https://example.org/payment")

        assert _content(docx_bytes, build) == (
            "[Claim form](https://example.org/claim)"
            "[Payment dates](https://example.org/payment)"
        )

    def test_a_bookmark_is_kept_exactly_as_word_wrote_it(self, docx_bytes):
        """Resolving one to a section is a separate question the PoC got wrong."""

        def build(document):
            _hyperlink(document.add_paragraph(), "Annex A", anchor="_Toc178312")

        assert _content(docx_bytes, build) == "[Annex A](#_Toc178312)"

    def test_a_target_that_would_not_parse_bare_is_bracketed(self, docx_bytes):
        """A space would end the destination, and a bracket would close it early."""

        def build(document):
            _hyperlink(
                document.add_paragraph(), "Claim form", url="https://example.org/a b(1)"
            )

        assert _content(docx_bytes, build) == (
            "[Claim form](<https://example.org/a b(1)>)"
        )

    def test_a_target_whose_brackets_balance_is_left_bare(self, docx_bytes):
        """CommonMark reads a balanced pair as part of the destination, and an
        address carrying brackets at all - a filename in a path - balances them. The
        editor writes such a target bare, so wrapping it here would have the first
        save rewrite the link, and any table holding it re-measured and re-padded."""

        def build(document):
            _hyperlink(
                document.add_paragraph(),
                "Claim form",
                url="https://example.org/a(1).pdf",
            )

        assert _content(docx_bytes, build) == (
            "[Claim form](https://example.org/a(1).pdf)"
        )

    def test_a_target_whose_brackets_do_not_balance_is_bracketed(self, docx_bytes):
        """A close with no open ends the destination where it sits, taking the rest
        of the address out of the link with it."""

        def build(document):
            _hyperlink(
                document.add_paragraph(), "Claim form", url="https://example.org/a)b"
            )

        assert _content(docx_bytes, build) == (
            "[Claim form](<https://example.org/a)b>)"
        )

    def test_the_styling_word_paints_on_a_link_is_not_repeated_in_the_text(
        self, docx_bytes
    ):
        """Every hyperlink is underlined and blue by style, so neither says anything.

        Dropping the colour here is also what keeps a coloured span and a link from
        ever nesting, which nothing reading the span's brackets back could unpick."""

        def build(document):
            runs = _hyperlink(
                document.add_paragraph(), "Guidance", url="https://example.org/guidance"
            )
            runs[0].font.underline = True
            _colour(runs[0], "0000FF")

        assert _content(docx_bytes, build) == "[Guidance](https://example.org/guidance)"

    def test_a_mark_of_the_authors_own_survives_a_link(self, docx_bytes):
        def build(document):
            runs = _hyperlink(
                document.add_paragraph(), "Guidance", url="https://example.org/guidance"
            )
            runs[0].font.bold = True

        assert _content(docx_bytes, build) == (
            "[**Guidance**](https://example.org/guidance)"
        )

    @pytest.mark.parametrize(
        "texts",
        [
            pytest.param([], id="no-runs-at-all"),
            pytest.param(["  "], id="nothing-but-space"),
        ],
    )
    def test_a_link_with_no_text_is_left_out(self, docx_bytes, texts):
        """An empty link renders as `[]()`, which is a defect on the page.

        Word writes a run-less w:hyperlink around a bookmark that its text has since
        been deleted from, so there is nothing to take the link's marks from either.
        """

        def build(document):
            paragraph = document.add_paragraph("See the guidance. ")
            _hyperlink(paragraph, *texts, url="https://example.org/guidance")

        assert _content(docx_bytes, build) == "See the guidance."

    def test_a_link_pointing_nowhere_is_left_as_its_text(self, docx_bytes):
        """A w:hyperlink with neither a relationship nor a bookmark has no target."""

        def build(document):
            _hyperlink(document.add_paragraph(), "Guidance")

        assert _content(docx_bytes, build) == "Guidance"


class TestFieldLinks:
    """Word's older HYPERLINK field, which writes a link without a w:hyperlink.

    Six of these across the two real guides reached the page as unlinked text: the
    address is in the field's instruction, and an instruction is not w:t.
    """

    def test_a_field_hyperlink_is_a_link(self, docx_bytes):
        def build(document):
            _field(
                document.add_paragraph(),
                ' HYPERLINK "https://example.org/claim" ',
                result="Claim form",
            )

        assert _content(docx_bytes, build) == (
            "[Claim form](https://example.org/claim)"
        )

    def test_an_instruction_split_across_runs_is_one_target(self, docx_bytes):
        """Word breaks a long instruction wherever it likes, mid-address included."""

        def build(document):
            _field(
                document.add_paragraph(),
                " HYPERLINK ",
                '"https://example.org/',
                'claim" ',
                result="Claim form",
            )

        assert _content(docx_bytes, build) == (
            "[Claim form](https://example.org/claim)"
        )

    def test_a_field_word_never_rendered_says_nothing(self, docx_bytes):
        """With no separate there is no result, and the instruction is not text."""

        def build(document):
            paragraph = document.add_paragraph("See the guidance. ")
            _field(paragraph, ' HYPERLINK "https://example.org/claim" ')

        assert _content(docx_bytes, build) == "See the guidance."

    def test_a_field_hyperlink_with_no_text_is_not_a_link(self, docx_bytes):
        """An empty link renders as `[]()`, which is a defect on the page."""

        def build(document):
            paragraph = document.add_paragraph("See the guidance. ")
            _field(paragraph, ' HYPERLINK "https://example.org/claim" ', result="   ")

        assert _content(docx_bytes, build) == "See the guidance."

    def test_a_field_left_open_still_puts_its_text_on_the_page(self, docx_bytes):
        """Word ends a paragraph between a field's text and its end boundary.

        It does so once in one of the two real guides, and waiting for an end that
        never comes drops the words along with the link.
        """

        def build(document):
            _field(
                document.add_paragraph(),
                ' HYPERLINK "https://example.org/claim" ',
                result="Claim form",
                closed=False,
            )

        assert _content(docx_bytes, build) == (
            "[Claim form](https://example.org/claim)"
        )

    def test_a_field_that_is_not_a_link_is_left_as_its_text(self, docx_bytes):
        """A page number is what PAGEREF renders; the bookmark it read is not text."""

        def build(document):
            _field(
                document.add_paragraph(),
                " PAGEREF _Ref17831 \\h ",
                result="12",
            )

        assert _content(docx_bytes, build) == "12"


class TestTabsAndBreaks:
    def test_a_tab_becomes_the_separator_it_is(self, docx_bytes):
        """Markdown has no tab stop, and a line opening with one is a code block."""

        def build(document):
            run = document.add_paragraph().add_run()
            run.add_tab()
            run.add_text("Parcel")
            run.add_tab()
            run.add_text("Area")

        assert _content(docx_bytes, build) == "Parcel Area"

    def test_a_soft_break_becomes_a_line_break_in_markdown(self, docx_bytes):
        """A break with no type wraps the line; joining the two loses the wrap."""

        def build(document):
            run = document.add_paragraph().add_run("Example Agency")
            _break(run)
            run.add_text("Claims team")

        assert _content(docx_bytes, build) == "Example Agency\\\nClaims team"

    def test_a_break_with_nothing_after_it_leaves_no_dangling_line(self, docx_bytes):
        """A trailing backslash with no line under it is a defect on the page."""

        def build(document):
            run = document.add_paragraph().add_run("Example Agency")
            _break(run)

        assert _content(docx_bytes, build) == "Example Agency"

    def test_a_page_break_says_nothing_about_the_text_around_it(self, docx_bytes):
        """Pagination is not punctuation: the sentence carries on over the page."""

        def build(document):
            paragraph = document.add_paragraph()
            run = paragraph.add_run("Apply before ")
            _break(run, "page")
            paragraph.add_run("the deadline.")

        assert _content(docx_bytes, build) == "Apply before the deadline."


class TestEscapingWhatTheAuthorTyped:
    """Text that happens to look like syntax has to be written so it says itself.

    The rule is the editor's own serialiser, character for character, so the
    Markdown the parser writes is already the Markdown the editor would save and
    the round trip has nothing left to change.
    """

    @pytest.mark.parametrize(
        ("typed", "written"),
        [
            pytest.param("<CS Claim Revenue>", "&lt;CS Claim Revenue&gt;", id="angles"),
            pytest.param(
                "Rural Payments & Land", "Rural Payments &amp; Land", id="amp"
            ),
        ],
    )
    def test_a_placeholder_is_not_handed_to_the_renderer_as_a_tag(
        self, docx_bytes, typed, written
    ):
        """23 runs across the two real guides read as HTML tags without this."""

        def build(document):
            document.add_paragraph().add_run(typed)

        assert _content(docx_bytes, build) == written

    def test_the_ampersand_is_encoded_before_the_angle_brackets(self, docx_bytes):
        """The other order gives `&amp;lt;`, which is the escape shown to the reader."""

        def build(document):
            document.add_paragraph().add_run("<a & b>")

        assert _content(docx_bytes, build) == "&lt;a &amp; b&gt;"

    @pytest.mark.parametrize(
        "character",
        [
            pytest.param("\\", id="backslash"),
            pytest.param("`", id="backtick"),
            pytest.param("*", id="asterisk"),
            pytest.param("_", id="underscore"),
            pytest.param("[", id="open-bracket"),
            pytest.param("]", id="close-bracket"),
            pytest.param("~", id="tilde"),
        ],
    )
    def test_each_syntax_character_is_written_as_itself(self, docx_bytes, character):
        """All seven unconditionally: which one would parse as syntax depends on
        where it sits, and over-escaping renders the same."""

        def build(document):
            document.add_paragraph().add_run(f"Field{character}name")

        assert _content(docx_bytes, build) == f"Field\\{character}name"

    def test_a_backslash_is_escaped_once_and_not_again_by_its_own_rule(
        self, docx_bytes
    ):
        """Escaping the escape character last would double every backslash added."""

        def build(document):
            document.add_paragraph().add_run("a\\b*c")

        assert _content(docx_bytes, build) == "a\\\\b\\*c"

    def test_an_asterisk_in_the_text_no_longer_opens_emphasis(self, docx_bytes):
        """The six asterisks of CS's filename rule, where `*[Title]*` disappeared."""

        def build(document):
            document.add_paragraph().add_run("[SBI]*[Title]*[Scheme Year]")

        assert _content(docx_bytes, build) == (
            "\\[SBI\\]\\*\\[Title\\]\\*\\[Scheme Year\\]"
        )

    def test_a_links_text_is_escaped_and_its_destination_is_not(self, docx_bytes):
        """The destination is an address, not prose: an escape in it breaks the link."""

        def build(document):
            _hyperlink(
                document.add_paragraph(),
                "Form_2 [draft]",
                url="https://example.org/a_b~c",
            )

        assert _content(docx_bytes, build) == (
            "[Form\\_2 \\[draft\\]](https://example.org/a_b~c)"
        )

    def test_the_markers_the_parser_writes_are_not_escaped(self, docx_bytes):
        """The rule sees the author's text only; every marker is added after it."""

        def build(document):
            paragraph = document.add_paragraph()
            paragraph.add_run("Note_1").bold = True
            paragraph.add_run(" ")
            _colour(paragraph.add_run("[urgent]"), "FF0000")

        assert _content(docx_bytes, build) == ("**Note\\_1** [\\[urgent\\]]{.red}")

    def test_a_hard_line_break_is_still_a_break_and_not_an_escaped_backslash(
        self, docx_bytes
    ):
        """The break's own backslash is written after the escape and stays one."""

        def build(document):
            run = document.add_paragraph().add_run("Team_A")
            _break(run)
            run.add_text("Team_B")

        assert _content(docx_bytes, build) == "Team\\_A\\\nTeam\\_B"
