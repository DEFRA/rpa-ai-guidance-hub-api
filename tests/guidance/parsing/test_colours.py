"""Reading the colour a run is written in, and writing it as Markdown.

Two things are being pinned here. The first is the matching rule: every colour a
document can carry has to come out as one of two names or as no colour at all, and the
cases below are the shapes that rule has to get right rather than a sample of them.

The second is the form. A coloured span is `[text]{.red}`, and what makes that work is
that its contents stay valid Markdown - so the cases that matter most are the ones where
the text holds something the syntax could otherwise trip over.

All fixture text is invented, as everywhere in this package.
"""

from collections.abc import Callable

import pytest
from docx.document import Document
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.text.run import Run

from app.guidance.parsing import parser

HEADING = "Applying"

# The `docx_bytes` fixture's contract, restated as it is in the sibling suites. It
# cannot be shared through conftest.py: pytest runs with --import-mode=importlib, so the
# test directory is never put on sys.path and `from conftest import ...` fails at
# collection.
Build = Callable[[Document], None]
DocxBytes = Callable[[Build], bytes]


def _colour_content(docx_bytes: DocxBytes, build: Build) -> str:
    """The Markdown of the section `build` writes its paragraphs into.

    Named for this file rather than `_content`, which sibling test modules define with
    signatures of their own; the package carries no __init__.py, so an editor indexing
    the directory has no boundary to tell it they are different functions.
    """

    def with_heading(document: Document) -> None:
        document.add_heading(HEADING, level=1)
        build(document)

    return parser.parse_docx(docx_bytes(with_heading)).sections[0].content


def _painted(run: Run, value: str) -> None:
    """Colour a run by raw w:val, including values python-docx will not write."""
    element = OxmlElement("w:color")
    element.set(qn("w:val"), value)
    run._r.get_or_add_rPr().append(element)


def _coloured(document: Document, text: str, value: str) -> None:
    """A paragraph of one coloured run."""
    _painted(document.add_paragraph().add_run(text), value)


class TestMatchingAColour:
    @pytest.mark.parametrize(
        ("value", "name"),
        [
            pytest.param("FF0000", "red", id="pure-red"),
            pytest.param("EE0000", "red", id="a-shade-off-red"),
            pytest.param("D13438", "red", id="a-red-with-blue-in-it"),
            pytest.param("0000FF", "blue", id="pure-blue"),
            pytest.param("0000E1", "blue", id="a-shade-off-blue"),
            pytest.param("0070C0", "blue", id="a-blue-with-green-in-it"),
            pytest.param("3333FF", "blue", id="a-blue-with-red-in-it"),
        ],
    )
    def test_a_colour_is_read_as_the_nearer_of_the_two(self, docx_bytes, value, name):
        """An author picking a shade by hand means the intent it approximates, so a
        one-off is read as the convention it was reaching for rather than kept as a
        stray nothing downstream would know what to do with."""

        def build(document):
            _coloured(document, "Mandatory", value)

        assert _colour_content(docx_bytes, build) == f"[Mandatory]{{.{name}}}"

    @pytest.mark.parametrize(
        "value",
        [
            pytest.param("000000", id="black"),
            pytest.param("808080", id="grey"),
            pytest.param("FFFFFF", id="white"),
            pytest.param("auto", id="auto"),
            pytest.param("", id="nothing-at-all"),
            pytest.param("FF00", id="too-short-to-be-a-colour"),
            pytest.param("ZZ0000", id="not-hex"),
        ],
    )
    def test_a_colour_naming_neither_is_no_colour(self, docx_bytes, value):
        """Word writes the default text colour out explicitly rather than leaving it
        unsaid, so most coloured runs in a document are wearing no colour at all."""

        def build(document):
            _coloured(document, "Ordinary", value)

        assert _colour_content(docx_bytes, build) == "Ordinary"

    @pytest.mark.parametrize(
        "value",
        [
            pytest.param("00FF00", id="green"),
            pytest.param("FF00FF", id="magenta"),
        ],
    )
    def test_a_colour_equally_far_from_both_names_neither(self, docx_bytes, value):
        """Equal red and blue channels put a colour on the plane between the two, so
        there is no nearer one to name. The same rule that drops the greys, which is
        why neither needs a case of its own in the matching."""

        def build(document):
            _coloured(document, "Ordinary", value)

        assert _colour_content(docx_bytes, build) == "Ordinary"


class TestWritingTheSpan:
    def test_marks_inside_a_coloured_span_stay_markdown(self, docx_bytes):
        """The editor re-tokenises a span's contents, so what goes inside one has to be
        Markdown. Emphasis written outside the span would be emphasis of the brackets."""

        def build(document):
            run = document.add_paragraph().add_run("Overdue")
            run.bold = True
            _painted(run, "FF0000")

        assert _colour_content(docx_bytes, build) == "[**Overdue**]{.red}"

    def test_a_bracket_in_the_text_is_escaped_inside_the_span(self, docx_bytes):
        """A bracket the author typed would close the span early, and whatever read it
        back would find no span there and drop the colour. Escaping is what stops that,
        so the placeholder conventions these documents are written in survive."""

        def build(document):
            _coloured(document, "[SBI]", "FF0000")

        assert _colour_content(docx_bytes, build) == "[\\[SBI\\]]{.red}"

    def test_two_shades_of_one_intent_are_a_single_span(self, docx_bytes):
        """Matching to a name before the runs are merged is what does this. Comparing
        the hex instead would break a phrase in half at a boundary no reader can see."""

        def build(document):
            paragraph = document.add_paragraph()
            _painted(paragraph.add_run("Complete "), "FF0000")
            _painted(paragraph.add_run("this part"), "EE0000")

        assert _colour_content(docx_bytes, build) == "[Complete this part]{.red}"

    def test_the_two_colours_are_two_spans(self, docx_bytes):
        """They say different things, so neighbouring runs do not merge across them."""

        def build(document):
            paragraph = document.add_paragraph()
            _painted(paragraph.add_run("Fill in"), "FF0000")
            _painted(paragraph.add_run("choose one"), "0000FF")

        assert (
            _colour_content(docx_bytes, build) == "[Fill in]{.red}[choose one]{.blue}"
        )

    def test_the_space_around_a_span_stays_outside_it(self, docx_bytes):
        """As for every other mark: a span whose brackets wrap a trailing space would
        put the colour on the gap between two words."""

        def build(document):
            paragraph = document.add_paragraph()
            _painted(paragraph.add_run("Complete "), "FF0000")
            paragraph.add_run("the note")

        assert _colour_content(docx_bytes, build) == "[Complete]{.red} the note"

    def test_a_coloured_run_saying_nothing_writes_no_span(self, docx_bytes):
        """Word leaves colour on runs holding no text - a paragraph mark, or an edit
        that removed the words but not their formatting. An empty span would be
        punctuation the reader has to step over."""

        def build(document):
            paragraph = document.add_paragraph()
            _painted(paragraph.add_run(""), "FF0000")
            paragraph.add_run("Send the form")

        assert _colour_content(docx_bytes, build) == "Send the form"
