"""Reading the boxes Word draws inside a run rather than beside it.

A text box is not a block of the body: it hangs off a w:drawing on a run, so no walk
of the body's own children reaches it. These cases are about getting to it, about
getting to it exactly once when Word has written it twice, and about it landing in
the section the page shows it in.

All fixture text is invented, as everywhere in this package.
"""

from collections.abc import Callable
from typing import Any

from docx.document import Document
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.text.paragraph import Paragraph
from lxml import etree

from app.guidance.parsing import parser

HEADING = "Applying"

# Markup compatibility, built through lxml rather than OxmlElement: python-docx's
# namespace map has no "mc" prefix to resolve, which is the same reason the parser
# spells these tags out in full.
_MC = "{http://schemas.openxmlformats.org/markup-compatibility/2006}"

# The `docx_bytes` fixture's contract, restated here as it is in test_inline.py. It
# cannot be shared through conftest.py: pytest runs with --import-mode=importlib, so
# the test directory is never put on sys.path and `from conftest import ...` fails at
# collection. Naming it is what stops an editor inferring these helpers' parameters
# structurally and guessing wrong.
Build = Callable[[Document], None]
DocxBytes = Callable[[Build], bytes]


def _box_content(docx_bytes: DocxBytes, build: Build, index: int = 0) -> str:
    """The Markdown of the section `build` writes its paragraphs into.

    Named for this file rather than `_content`, which two sibling test modules also
    define with a signature of their own. Nothing at runtime confuses them - pytest
    imports each file as its own module - but the package carries no __init__.py, so
    an editor indexing the directory has no boundary to tell it that three same-named
    helpers are three different functions.
    """

    def with_heading(document: Document) -> None:
        document.add_heading(HEADING, level=1)
        build(document)

    return parser.parse_docx(docx_bytes(with_heading)).sections[index].content


def _text_box(*texts: str) -> Any:
    """A w:txbxContent holding one paragraph per text given."""
    box = OxmlElement("w:txbxContent")
    for text in texts:
        paragraph = OxmlElement("w:p")
        run = OxmlElement("w:r")
        element = OxmlElement("w:t")
        element.text = text
        run.append(element)
        paragraph.append(run)
        box.append(paragraph)
    return box


def _anchor(paragraph: Paragraph, box: Any) -> None:
    """Hang a text box off a paragraph the way a drawing does.

    Only the nesting matters to the walk, not the shape around it, so the w:drawing
    is written without the DrawingML that would tell Word how big to draw it.
    """
    run = paragraph.add_run()
    drawing = OxmlElement("w:drawing")
    drawing.append(box)
    run._r.append(drawing)


def _alternate(
    paragraph: Paragraph, *, choice: Any = None, fallback: Any = None
) -> Any:
    """Hang an mc:AlternateContent off a paragraph, holding only the branches given.

    A branch is written only where a box is passed for it, because an empty
    mc:Choice is still a choice: Word would take it and draw nothing, which is not
    what a test naming only a fallback means to describe.
    """
    alternate = etree.Element(f"{_MC}AlternateContent")
    for tag, box in ((f"{_MC}Choice", choice), (f"{_MC}Fallback", fallback)):
        if box is None:
            continue
        drawing = OxmlElement("w:drawing")
        drawing.append(box)
        etree.SubElement(alternate, tag).append(drawing)

    paragraph.add_run()._r.append(alternate)
    return alternate


class TestReachingTheBox:
    def test_a_text_box_becomes_a_blockquote(self, docx_bytes):
        """Word draws a box round it, exactly as it does round a one-cell table, so
        the two say the same thing in Markdown."""

        def build(document):
            _anchor(document.add_paragraph(), _text_box("Check the version"))

        assert _box_content(docx_bytes, build) == "> Check the version"

    def test_a_box_keeps_its_paragraphs_apart(self, docx_bytes):
        """Reading the box as one string is what put a case-note template's five
        paragraphs on a single line."""

        def build(document):
            _anchor(document.add_paragraph(), _text_box("Version used:", "Name:"))

        assert _box_content(docx_bytes, build) == "> Version used:\n>\n> Name:"

    def test_the_anchoring_paragraph_still_says_what_it_says(self, docx_bytes):
        """The box is extra, not instead: a paragraph carrying one is a paragraph."""

        def build(document):
            paragraph = document.add_paragraph("Use the following template.")
            _anchor(paragraph, _text_box("Check the version"))

        assert _box_content(docx_bytes, build) == (
            "Use the following template.\n\n> Check the version"
        )

    def test_marks_inside_a_box_survive(self, docx_bytes):
        """A box's paragraphs are read as paragraphs, not as text, so everything an
        ordinary paragraph keeps is kept here too."""

        def build(document):
            box = _text_box("")
            run = box.findall(qn("w:p"))[0].findall(qn("w:r"))[0]
            properties = OxmlElement("w:rPr")
            properties.append(OxmlElement("w:b"))
            run.insert(0, properties)
            run.findall(qn("w:t"))[0].text = "Withdrawal"
            _anchor(document.add_paragraph(), box)

        assert _box_content(docx_bytes, build) == "> **Withdrawal**"


class TestWordWritingItTwice:
    def test_only_the_choice_is_read_when_both_branches_are_written(self, docx_bytes):
        """Word writes one copy of a shape per consumer that might render it and
        draws exactly one. Taking both said a real guide's box twice over - 85 words
        of it - and the audit counting the same pair is what hid it."""

        def build(document):
            _alternate(
                document.add_paragraph(),
                choice=_text_box("Check the version"),
                fallback=_text_box("Check the version"),
            )

        assert _box_content(docx_bytes, build) == "> Check the version"

    def test_the_fallback_is_read_when_there_is_no_choice(self, docx_bytes):
        """A branch is only unread because a better one was taken, never because it
        is a fallback."""

        def build(document):
            _alternate(document.add_paragraph(), fallback=_text_box("Check the claim"))

        assert _box_content(docx_bytes, build) == "> Check the claim"

    def test_an_alternate_offering_nothing_says_nothing(self, docx_bytes):
        def build(document):
            _alternate(document.add_paragraph())
            document.add_paragraph("Send the form")

        assert _box_content(docx_bytes, build) == "Send the form"

    def test_a_box_inside_a_box_is_not_said_again(self, docx_bytes):
        """The inner box is the outer one's own content. Emitting it here as well
        would repeat whatever the outer box already printed."""

        def build(document):
            outer = _text_box("Check the version")
            _anchor(document.add_paragraph(), outer)
            inner = _text_box("Check the claim")
            drawing = OxmlElement("w:drawing")
            drawing.append(inner)
            outer.findall(qn("w:p"))[0].findall(qn("w:r"))[0].append(drawing)

        assert _box_content(docx_bytes, build) == "> Check the version"


class TestWhereABoxGoes:
    def test_a_box_files_under_the_section_open_at_the_time(self, docx_bytes):
        def build(document):
            _anchor(document.add_paragraph(), _text_box("Check the version"))
            document.add_heading("Assessing", level=1)
            _anchor(document.add_paragraph(), _text_box("Check the claim"))

        sections = parser.parse_docx(
            docx_bytes(lambda d: (d.add_heading(HEADING, level=1), build(d)))
        ).sections
        assert [section.content for section in sections] == [
            "> Check the version",
            "> Check the claim",
        ]

    def test_a_box_on_a_heading_belongs_to_that_heading(self, docx_bytes):
        """The box is yielded after the paragraph it hangs from, so a heading has
        opened its section by the time its own box arrives."""

        def build(document):
            heading = document.add_heading("Assessing", level=1)
            _anchor(heading, _text_box("Check the claim"))

        assert _box_content(docx_bytes, build, index=1) == "> Check the claim"

    def test_a_box_closes_an_open_list_run(self, docx_bytes):
        """Two runs either side of a box are two lists, as they are either side of
        a table."""

        def build(document):
            document.add_paragraph("Send the form", style="List Bullet")
            _anchor(document.add_paragraph(), _text_box("Check the version"))
            document.add_paragraph("Keep a copy", style="List Bullet")

        assert _box_content(docx_bytes, build) == (
            "- Send the form\n\n> Check the version\n\n- Keep a copy"
        )

    def test_a_box_ahead_of_every_heading_is_dropped(self, docx_bytes):
        """It has no section to belong to, exactly as a paragraph there has none."""

        def build(document):
            _anchor(document.add_paragraph(), _text_box("Check the version"))

        assert parser.parse_docx(docx_bytes(build)).sections == []

    def test_an_empty_box_says_nothing(self, docx_bytes):
        def build(document):
            _anchor(document.add_paragraph(), _text_box())
            document.add_paragraph("Send the form")

        assert _box_content(docx_bytes, build) == "Send the form"
