"""Shared fixtures for the parsing tests.

All fixture text in this package is invented. No real guidance document, scheme or
organisation name appears anywhere in the suite - keep it that way when adding cases.
"""

import io

import docx
import pytest
from docx.enum.style import WD_STYLE_TYPE


@pytest.fixture
def docx_bytes():
    """Return a builder for an in-memory .docx.

    Takes a callable that populates the document, so a test can use the whole of
    python-docx's API rather than whatever vocabulary a declarative fixture would
    have to invent. `core_title` sets the docProps title, which is separate from
    anything printed on the page.
    """

    def build_docx(build=None, core_title=None) -> bytes:
        document = docx.Document()
        if build is not None:
            build(document)
        if core_title is not None:
            document.core_properties.title = core_title

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    return build_docx


@pytest.fixture
def in_style():
    """Return a helper adding a paragraph in a named style.

    python-docx will only apply a style the template already defines, and the
    styles worth testing against - a custom annex style, a heading-like name that
    is not a heading - are exactly the ones it does not have.
    """

    def add_paragraph(document, text: str, style_name: str) -> None:
        if all(style.name != style_name for style in document.styles):
            document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        document.add_paragraph(text, style=style_name)

    return add_paragraph
