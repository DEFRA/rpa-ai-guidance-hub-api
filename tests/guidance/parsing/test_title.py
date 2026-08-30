"""Inferring the document's title.

The cover page is read first and the docProps title is only a fallback, so most of
these cases are about reconstructing what the cover actually printed, and about
knowing where the cover stops.

Where a case is about the cover *ending*, the text after the boundary carries the
Title style too. Without that, the "Title-styled paragraphs win" filter in
`_cover_title` discards that text whether or not the boundary was found, and the
test passes however badly the boundary logic is broken.
"""

import pytest
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn

from app.guidance.parsing import parser

TITLE = "Example Grant Scheme Guide"
LONG_TITLE = "Example Grant Scheme Processing to Final Payment Guide"
OVERLEAF = "Not part of the title"


def _page_break(document) -> None:
    """Add the paragraph Word inserts for an explicit page break."""
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _in_style(document, text: str, style_name: str) -> None:
    """Add a paragraph in a style, defining it first if the template lacks it."""
    if all(style.name != style_name for style in document.styles):
        document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    document.add_paragraph(text, style=style_name)


class TestReconstructingTheCoverTitle:
    def test_title_styled_paragraphs_win_over_other_cover_text(self, docx_bytes):
        def build(document):
            document.add_paragraph("Example Agency")
            document.add_paragraph("")
            document.add_paragraph(TITLE, style="Title")
            _page_break(document)

        assert parser.parse_docx(docx_bytes(build)).title == TITLE

    def test_hand_wrapped_lines_are_rejoined_into_one_part(self, docx_bytes):
        """Consecutive lines are one title the author wrapped, not two parts.

        Each line is padded, because a hand-wrapped one so often is.
        """

        def build(document):
            document.add_paragraph("Example Grant Scheme Processing  ", style="Title")
            document.add_paragraph("  to Final Payment Guide", style="Title")
            _page_break(document)

        assert parser.parse_docx(docx_bytes(build)).title == LONG_TITLE

    def test_a_blank_line_separates_two_parts_of_a_title(self, docx_bytes):
        def build(document):
            document.add_paragraph(TITLE, style="Title")
            document.add_paragraph("")
            document.add_paragraph("2026 edition", style="Title")
            _page_break(document)

        assert parser.parse_docx(docx_bytes(build)).title == f"{TITLE} — 2026 edition"

    def test_padding_blank_paragraphs_do_not_form_parts_of_their_own(self, docx_bytes):
        """No paragraph here carries the Title style, deliberately.

        The style filter would otherwise drop the empty groups on its way past, and
        the case would pass whether or not runs of blanks are handled.
        """

        def build(document):
            document.add_paragraph("")
            document.add_paragraph("")
            document.add_paragraph("Example Agency")
            document.add_paragraph("")
            document.add_paragraph(TITLE)
            document.add_paragraph("")
            _page_break(document)

        assert parser.parse_docx(docx_bytes(build)).title == f"Example Agency — {TITLE}"

    def test_a_style_with_no_name_is_ordinary_cover_text(self, docx_bytes):
        """A style carrying no w:name element has no name python-docx can report.

        The paragraph is then neither Title-styled nor a boundary, so it is read as
        the plain cover text it is rather than bringing the parse down.
        """

        def build(document):
            document.styles.add_style("Nameless", WD_STYLE_TYPE.PARAGRAPH)
            document.add_paragraph(TITLE, style="Nameless")
            nameless = document.styles["Nameless"]._element
            nameless.remove(nameless.find(qn("w:name")))
            _page_break(document)

        assert parser.parse_docx(docx_bytes(build)).title == TITLE


class TestWhereTheCoverEnds:
    def test_a_page_break_ends_the_cover(self, docx_bytes):
        def build(document):
            document.add_paragraph(TITLE, style="Title")
            _page_break(document)
            document.add_paragraph(OVERLEAF, style="Title")

        assert parser.parse_docx(docx_bytes(build)).title == TITLE

    def test_a_paragraph_carrying_a_trailing_break_is_still_cover_text(
        self, docx_bytes
    ):
        """The break ends the page *after* this paragraph, so its words are cover."""

        def build(document):
            paragraph = document.add_paragraph(TITLE, style="Title")
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            document.add_paragraph(OVERLEAF, style="Title")

        assert parser.parse_docx(docx_bytes(build)).title == TITLE

    def test_a_line_break_does_not_end_the_cover(self, docx_bytes):
        """Only a page break ends the page. A line break is how a title wraps."""

        def build(document):
            paragraph = document.add_paragraph(TITLE, style="Title")
            paragraph.add_run().add_break(WD_BREAK.LINE)
            document.add_paragraph("2026 edition", style="Title")
            _page_break(document)

        assert parser.parse_docx(docx_bytes(build)).title == f"{TITLE} 2026 edition"

    def test_a_paragraph_forced_onto_a_new_page_is_not_cover_text(self, docx_bytes):
        """pageBreakBefore means the opposite: this paragraph is already overleaf."""

        def build(document):
            document.add_paragraph(TITLE, style="Title")
            overleaf = document.add_paragraph(OVERLEAF, style="Title")
            overleaf.paragraph_format.page_break_before = True

        assert parser.parse_docx(docx_bytes(build)).title == TITLE

    def test_a_disabled_page_break_before_leaves_a_paragraph_on_the_cover(
        self, docx_bytes
    ):
        """Turning the toggle off writes w:val="0" rather than removing it.

        Presence alone would therefore read as 'on' and lose the rest of the cover.
        """

        def build(document):
            document.add_paragraph(TITLE, style="Title")
            following = document.add_paragraph("2026 edition", style="Title")
            following.paragraph_format.page_break_before = False
            _page_break(document)

        assert parser.parse_docx(docx_bytes(build)).title == f"{TITLE} 2026 edition"

    @pytest.mark.parametrize(
        "style_name",
        [
            pytest.param("TOC Heading", id="toc-heading"),
            pytest.param("Contents 1", id="contents"),
            pytest.param("Table of Contents Body", id="table-of-contents"),
        ],
    )
    def test_a_table_of_contents_ends_the_cover(self, docx_bytes, style_name):
        """Word spells its navigation styles more than one way."""

        def build(document):
            document.add_paragraph(TITLE, style="Title")
            _in_style(document, "1 Introduction ....... 3", style_name)
            document.add_paragraph(OVERLEAF, style="Title")

        assert parser.parse_docx(docx_bytes(build)).title == TITLE

    def test_a_body_heading_ends_the_cover_with_no_break_at_all(self, docx_bytes):
        """Without this a coverless document swallows its opening heading."""

        def build(document):
            document.add_paragraph(TITLE, style="Title")
            document.add_paragraph("1 Introduction", style="Heading 1")
            document.add_paragraph(OVERLEAF, style="Title")

        assert parser.parse_docx(docx_bytes(build)).title == TITLE

    def test_an_appendix_heading_ends_the_cover(self, docx_bytes):
        """Word's own annex style is not a "Heading n", so it is matched by name."""

        def build(document):
            document.add_paragraph(TITLE, style="Title")
            _in_style(document, "Appendix A", "Appendix Heading")
            document.add_paragraph(OVERLEAF, style="Title")

        assert parser.parse_docx(docx_bytes(build)).title == TITLE


class TestFallingBackToStoredProperties:
    def test_core_properties_are_used_when_there_is_no_cover(self, docx_bytes):
        """A stored title is padded as often as not, so it is stripped.

        The PoC ignored one hardcoded template name so the cover could win. Reading
        the cover first makes that workaround unnecessary: where there is no cover at
        all, the stored name is the only title the document has, whatever it says.
        """

        def build(document):
            document.add_paragraph("1 Introduction", style="Heading 1")

        document = parser.parse_docx(docx_bytes(build, core_title="  Stored Title  "))

        assert document.title == "Stored Title"

    def test_the_printed_cover_beats_stale_core_properties(self, docx_bytes):
        """The real case: properties carried forward from the file this was copied from.

        Core properties said "2024", the cover printed no year, and the file was the
        2026 edition. Metadata goes stale silently; the cover is what a reader sees.
        """

        def build(document):
            document.add_paragraph(LONG_TITLE, style="Title")
            _page_break(document)

        document = parser.parse_docx(docx_bytes(build, core_title=f"2024 {LONG_TITLE}"))

        assert document.title == LONG_TITLE

    def test_a_document_with_neither_has_an_empty_title(self, docx_bytes):
        def build(document):
            document.add_paragraph("1 Introduction", style="Heading 1")

        assert parser.parse_docx(docx_bytes(build)).title == ""
