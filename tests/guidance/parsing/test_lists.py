"""Turning the numbering Word attaches to a paragraph into a Markdown list.

Word writes a list item as an ordinary paragraph carrying a w:numPr, and generates
the bullet or the digit only when it renders the page. Neither is in the text, so an
ordered list arrives with its ordinals missing entirely and these cases are what puts
them back.

All fixture text is invented, as everywhere in this package. The default template
already defines "List Bullet" (bullet) and "List Number" (decimal), but declares them
at ilvl 0 and nowhere else, so anything about a deeper level has to build a list
definition of its own.
"""

import pytest
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

from app.guidance.parsing import parser


def _val(name: str, value: object):
    """An element carrying nothing but a w:val, as OOXML writes most settings."""
    element = OxmlElement(name)
    element.set(qn("w:val"), str(value))
    return element


def _numbered(paragraph, num_id: int, level: int | None = None) -> None:
    """Attach numbering to a paragraph directly, over whatever its style says."""
    properties = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    if level is not None:
        properties.get_or_add_ilvl().set(qn("w:val"), str(level))
    properties.get_or_add_numId().set(qn("w:val"), str(num_id))


def _numbering_style(document, name: str, num_id: int) -> str:
    """Add a paragraph style that declares numbering, as a real template does.

    The name is deliberately the caller's to choose: what makes the paragraphs in
    this style list items is the numbering it declares, never what it is called.
    """
    style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    properties = style.element.get_or_add_pPr().get_or_add_numPr()
    properties.get_or_add_numId().set(qn("w:val"), str(num_id))
    return name


def _list_definition(document, *formats: str) -> int:
    """Add a list defining one format per level, and return the numId naming it."""
    numbering = document.part.numbering_part.element
    definitions = numbering.findall(qn("w:abstractNum"))
    abstract_id = 1 + max(
        int(definition.get(qn("w:abstractNumId"))) for definition in definitions
    )
    num_id = 1 + max(
        int(num.get(qn("w:numId"))) for num in numbering.findall(qn("w:num"))
    )

    definition = OxmlElement("w:abstractNum")
    definition.set(qn("w:abstractNumId"), str(abstract_id))
    for level, number_format in enumerate(formats):
        element = OxmlElement("w:lvl")
        element.set(qn("w:ilvl"), str(level))
        element.append(_val("w:numFmt", number_format))
        definition.append(element)
    definitions[-1].addnext(definition)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    num.append(_val("w:abstractNumId", abstract_id))
    numbering.append(num)
    return num_id


def _content(source: bytes, index: int = 0) -> str:
    """The Markdown one parsed section holds, in document order."""
    return parser.parse_docx(source).sections[index].content


@pytest.fixture
def opened(docx_bytes):
    """Build a document that already has a section open for the items to land in."""

    def build_docx(populate) -> bytes:
        def build(document):
            document.add_heading("Applying", level=1)
            populate(document)

        return docx_bytes(build)

    return build_docx


class TestFindingListItems:
    def test_a_run_of_numbered_paragraphs_is_one_block(self, opened):
        """One list, not one list per item: the run is what the reader sees."""

        def build(document):
            document.add_paragraph("Send the form", style="List Bullet")
            document.add_paragraph("Keep a copy", style="List Bullet")

        assert _content(opened(build)) == "- Send the form\n- Keep a copy"

    def test_an_ordered_list_carries_the_ordinals_word_generates(self, opened):
        """Word puts the digits on the page, never in the text, so they are lost
        entirely without this: one real guide's list says "go to step 8"."""

        def build(document):
            for step in ("Open the register", "Filter by reference", "Read the status"):
                document.add_paragraph(step, style="List Number")

        assert _content(opened(build)) == (
            "1. Open the register\n2. Filter by reference\n3. Read the status"
        )

    def test_numbering_removed_from_a_paragraph_leaves_it_as_prose(self, opened):
        """numId 0 does not name a list. It says numbering has been taken off this
        paragraph, and it has to beat the style, which still declares some."""

        def build(document):
            _numbered(
                document.add_paragraph("Before you start", style="List Bullet"), 0
            )

        assert _content(opened(build)) == "Before you start"

    def test_a_style_whose_name_says_nothing_still_makes_items(self, opened):
        """What makes an item is the numbering the style declares, not its name: 64
        items across the two real guides are styled plain "Normal"."""

        def build(document):
            style = _numbering_style(document, "Question", num_id=1)
            document.add_paragraph("Is the claim valid", style=style)

        assert _content(opened(build)) == "- Is the claim valid"

    def test_a_style_naming_a_bullet_but_declaring_none_is_prose(self, opened):
        """The other half of the same rule, and what Word itself prints."""

        def build(document):
            document.styles.add_style("List Bullet 4", WD_STYLE_TYPE.PARAGRAPH)
            document.add_paragraph("Before you start", style="List Bullet 4")

        assert _content(opened(build)) == "Before you start"

    def test_a_numbered_heading_opens_a_section_rather_than_an_item(self, docx_bytes):
        """Both real guides attach numbering to their heading styles, so 46 headings
        would become bullets if the list question were asked before the heading."""

        def build(document):
            _numbered(document.add_heading("Applying", level=1), 1)
            document.add_paragraph("Send the form")

        document = parser.parse_docx(docx_bytes(build))
        assert [(s.number, s.heading) for s in document.sections] == [("1", "Applying")]
        assert document.sections[0].content == "Send the form"

    def test_a_numbered_contents_entry_is_not_an_item(self, opened, in_style):
        """Contents are navigation wherever they turn up, however Word numbers them."""

        def build(document):
            in_style(document, "Applying", "toc 1")
            _numbered(document.paragraphs[-1], 1)

        assert _content(opened(build)) == ""


class TestReadingTheMarkerFormat:
    def test_the_format_comes_from_the_level_the_item_sits_at(self, opened):
        def build(document):
            num_id = _list_definition(document, "bullet", "decimal")
            _numbered(document.add_paragraph("Check the register"), num_id, level=0)
            _numbered(document.add_paragraph("Note the reference"), num_id, level=1)

        assert (
            _content(opened(build)) == "- Check the register\n  1. Note the reference"
        )

    def test_a_level_the_list_never_declares_takes_its_first(self, opened):
        """The only thing the document says about the list is what it says at the
        levels it did define."""

        def build(document):
            num_id = _list_definition(document, "decimal")
            _numbered(document.add_paragraph("Note the reference"), num_id, level=1)

        assert _content(opened(build)) == "1. Note the reference"

    def test_a_list_the_document_never_declares_is_bulleted(self, opened):
        """A numId no w:num carries says nothing about its format, and neither does
        a document with no numbering definitions at all."""

        def build(document):
            _numbered(document.add_paragraph("Check the register"), 99)

        assert _content(opened(build)) == "- Check the register"


class TestNesting:
    def test_a_run_opening_at_a_deeper_level_still_starts_at_the_left(self, opened):
        """Levels are relative, exactly as headings are."""

        def build(document):
            _numbered(document.add_paragraph("Check the register"), 1, level=2)

        assert _content(opened(build)) == "- Check the register"

    def test_a_skipped_level_nests_one_deep(self, opened):
        def build(document):
            _numbered(document.add_paragraph("Check the register"), 1, level=0)
            _numbered(document.add_paragraph("Note the reference"), 1, level=2)

        assert _content(opened(build)) == "- Check the register\n  - Note the reference"

    def test_a_child_hangs_from_its_parent_own_text(self, opened):
        """An ordered marker is wider than a bullet, so a fixed indent per level
        would leave a bullet child of a numbered parent outside the list."""

        def build(document):
            num_id = _list_definition(document, "decimal", "bullet")
            _numbered(document.add_paragraph("Open the register"), num_id, level=0)
            _numbered(document.add_paragraph("Filter by reference"), num_id, level=1)

        assert _content(opened(build)) == (
            "1. Open the register\n   - Filter by reference"
        )

    def test_the_count_carries_on_at_the_level_returned_to(self, opened):
        """A sub-list interrupting an ordered list must not restart it at 1."""

        def build(document):
            num_id = _list_definition(document, "decimal", "bullet")
            _numbered(document.add_paragraph("Open the register"), num_id, level=0)
            _numbered(document.add_paragraph("Filter by reference"), num_id, level=1)
            _numbered(document.add_paragraph("Read the status"), num_id, level=0)

        assert _content(opened(build)).endswith("\n2. Read the status")

    def test_a_level_opened_again_restarts_its_count(self, opened):
        def build(document):
            num_id = _list_definition(document, "decimal", "decimal")
            _numbered(document.add_paragraph("Open the register"), num_id, level=0)
            _numbered(document.add_paragraph("Filter by reference"), num_id, level=1)
            _numbered(document.add_paragraph("Read the status"), num_id, level=0)
            _numbered(document.add_paragraph("Note the outcome"), num_id, level=1)

        assert _content(opened(build)) == (
            "1. Open the register\n"
            "   1. Filter by reference\n"
            "2. Read the status\n"
            "   1. Note the outcome"
        )

    def test_a_bullet_interrupting_an_ordered_list_starts_a_new_one(self, opened):
        """What the ordered list had counted to says nothing about the bullets that
        follow it, and a blank line is what says so: a change of kind at the
        outermost depth is a second list, and the editor writes one there."""

        def build(document):
            document.add_paragraph("Open the register", style="List Number")
            document.add_paragraph("Filter by reference", style="List Bullet")
            document.add_paragraph("Read the status", style="List Number")

        assert _content(opened(build)) == (
            "1. Open the register\n\n- Filter by reference\n\n1. Read the status"
        )

    def test_a_change_of_kind_deeper_in_is_not_parted(self, opened):
        """A blank line inside a list makes it loose, so the editor writes none for
        a nested change however plainly it is a second list."""

        def build(document):
            num_id = _list_definition(document, "bullet", "bullet")
            other = _list_definition(document, "bullet", "decimal")
            _numbered(document.add_paragraph("Open the register"), num_id, level=0)
            _numbered(document.add_paragraph("Filter by reference"), num_id, level=1)
            _numbered(document.add_paragraph("Note the outcome"), other, level=1)

        assert _content(opened(build)) == (
            "- Open the register\n  - Filter by reference\n  1. Note the outcome"
        )


class TestClosingTheRun:
    def test_an_empty_paragraph_does_not_close_the_run(self, opened):
        """Word spaces its lists with empty paragraphs. Closing on one splits a
        single list into two blocks, and the blank line between them makes it a
        *loose* list - which the editor rewrites as the one tight list it always
        was, so the converted document would be changed by its first save."""

        def build(document):
            document.add_paragraph("Open the register", style="List Number")
            document.add_paragraph("")
            document.add_paragraph("Read the status", style="List Number")

        assert _content(opened(build)) == ("1. Open the register\n2. Read the status")

    def test_a_paragraph_closes_the_run_and_the_next_starts_again(self, opened):
        def build(document):
            document.add_paragraph("Open the register", style="List Number")
            document.add_paragraph("Then:")
            document.add_paragraph("Read the status", style="List Number")

        assert _content(opened(build)) == (
            "1. Open the register\n\nThen:\n\n1. Read the status"
        )

    def test_a_heading_closes_the_run_into_the_section_it_opened_in(self, docx_bytes):
        def build(document):
            document.add_heading("Applying", level=1)
            document.add_paragraph("Send the form", style="List Bullet")
            document.add_heading("Assessing", level=1)
            document.add_paragraph("Read the status", style="List Bullet")

        sections = parser.parse_docx(docx_bytes(build)).sections
        assert [section.content for section in sections] == [
            "- Send the form",
            "- Read the status",
        ]

    def test_an_item_saying_nothing_adds_no_bullet(self, opened):
        """An empty numbered paragraph is a layout artefact, as an empty paragraph
        is, and does not spend an ordinal on the way past."""

        def build(document):
            document.add_paragraph("Open the register", style="List Number")
            document.add_paragraph("   ", style="List Number")
            document.add_paragraph("Read the status", style="List Number")

        assert _content(opened(build)) == "1. Open the register\n2. Read the status"

    def test_a_line_break_inside_an_item_does_not_close_the_list(self, opened):
        """A hard break left in the first column would end the list at it."""

        def build(document):
            paragraph = document.add_paragraph("Open the register", style="List Number")
            paragraph.add_run().add_break()
            paragraph.add_run("Then filter it")

        assert _content(opened(build)) == ("1. Open the register\\\n   Then filter it")
