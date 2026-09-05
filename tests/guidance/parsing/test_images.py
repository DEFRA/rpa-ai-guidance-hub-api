"""Reading the pictures a document draws.

An image is inline content - every one in the two real guides is a `wp:inline` - so
these cases are mostly about it staying where the run put it, and about the name it is
given, which only the section it lands in can decide.

All fixture text is invented, as everywhere in this package.
"""

import base64
import io
from collections.abc import Callable
from typing import Any

from docx.document import Document
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.text.paragraph import Paragraph

from app.guidance.parsing import parser

HEADING = "Applying"

# The smallest PNG that python-docx will take: one transparent pixel. It has to be a
# real image rather than any old bytes, because python-docx reads the header to find
# the dimensions it writes into the drawing.
PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNkYPhfDwAChwGA60e6kgAAAABJRU5ErkJggg=="
)

# The `docx_bytes` fixture's contract, restated as it is in test_inline.py. It cannot
# be shared through conftest.py: pytest runs with --import-mode=importlib, so the test
# directory is never put on sys.path and `from conftest import ...` fails at collection.
Build = Callable[[Document], None]
DocxBytes = Callable[[Build], bytes]


def _image_content(docx_bytes: DocxBytes, build: Build, index: int = 0) -> str:
    """The Markdown of the section `build` writes its paragraphs into.

    Named for this file rather than `_content`, which sibling test modules define with
    signatures of their own. Nothing at runtime confuses them, but the package carries
    no __init__.py, so an editor indexing the directory has no boundary to tell it
    that same-named helpers are different functions.
    """

    def with_heading(document: Document) -> None:
        document.add_heading(HEADING, level=1)
        build(document)

    return parser.parse_docx(docx_bytes(with_heading)).sections[index].content


def _picture(paragraph: Paragraph) -> Any:
    """Draw a picture in a run of its own, the way Word writes one."""
    return paragraph.add_run().add_picture(io.BytesIO(PNG))


def _empty_drawing(paragraph: Paragraph) -> None:
    """A drawing holding no picture, as a shape or a chart would be."""
    paragraph.add_run()._r.append(OxmlElement("w:drawing"))


def _unembed(paragraph: Paragraph) -> None:
    """Turn an embedded picture into a linked one, which carries no r:embed."""
    for blip in paragraph._p.iter(qn("a:blip")):
        del blip.attrib[qn("r:embed")]


class TestWhereAPictureGoes:
    def test_a_picture_becomes_an_image_named_for_its_section(self, docx_bytes):
        """Only a section can say what a picture is called, so the name is given
        after the walk rather than where the picture is met."""

        def build(document):
            _picture(document.add_paragraph())

        assert _image_content(docx_bytes, build) == "![](1_img_1.png)"

    def test_a_picture_beside_text_stays_where_it_sits(self, docx_bytes):
        """One of the two real inline icons reads "Select the (icon) binocular
        icon", so putting the picture on a line of its own would move it out of the
        sentence that names it."""

        def build(document):
            paragraph = document.add_paragraph()
            paragraph.add_run("Select the ")
            _picture(paragraph)
            paragraph.add_run(" binocular icon")

        assert _image_content(docx_bytes, build) == (
            "Select the ![](1_img_1.png) binocular icon"
        )

    def test_a_picture_in_a_list_item_keeps_its_marker(self, docx_bytes):
        def build(document):
            _picture(document.add_paragraph(style="List Bullet"))

        assert _image_content(docx_bytes, build) == "- ![](1_img_1.png)"

    def test_two_pictures_side_by_side_are_two_pictures(self, docx_bytes):
        """They merge with nothing: a run of text either side of them would be one
        span, but a picture is not text."""

        def build(document):
            paragraph = document.add_paragraph()
            _picture(paragraph)
            _picture(paragraph)

        assert _image_content(docx_bytes, build) == "![](1_img_1.png)![](1_img_2.png)"

    def test_a_picture_takes_no_marks(self, docx_bytes):
        """Five image runs in the two real guides are bold, struck through or
        coloured. That is Word's formatting reaching a place where it says nothing,
        and wrapping a picture in it would put markers round an image."""

        def build(document):
            paragraph = document.add_paragraph()
            run = paragraph.add_run()
            run.bold = True
            run.add_picture(io.BytesIO(PNG))

        assert _image_content(docx_bytes, build) == "![](1_img_1.png)"

    def test_a_drawing_holding_no_picture_says_nothing(self, docx_bytes):
        """A shape or a chart is a drawing too, and neither is an image."""

        def build(document):
            _empty_drawing(document.add_paragraph())
            document.add_paragraph("Send the form")

        assert _image_content(docx_bytes, build) == "Send the form"

    def test_a_linked_picture_says_nothing(self, docx_bytes):
        """A picture Word links to rather than embeds writes r:link and no r:embed:
        its bytes live outside the package, so there is nothing here to carry and
        nothing that could be written back out under a generated name."""

        def build(document):
            paragraph = document.add_paragraph()
            _picture(paragraph)
            _unembed(paragraph)
            document.add_paragraph("Send the form")

        assert _image_content(docx_bytes, build) == "Send the form"


class TestNaming:
    def test_pictures_are_numbered_within_their_own_section(self, docx_bytes):
        """The number says where the picture belongs, so it restarts wherever a new
        section does."""

        def build(document):
            _picture(document.add_paragraph())
            _picture(document.add_paragraph())
            document.add_heading("Assessing", level=1)
            _picture(document.add_paragraph())

        document = parser.parse_docx(
            docx_bytes(lambda d: (d.add_heading(HEADING, level=1), build(d)))
        )
        assert [image.name for image in document.images] == [
            "1_img_1.png",
            "1_img_2.png",
            "2_img_1.png",
        ]

    def test_a_picture_carries_the_bytes_and_type_of_its_part(self, docx_bytes):
        """The extension is the part's own, never assumed - one image in the two
        real guides is a JPEG among fifty-eight PNGs."""

        def build(document):
            _picture(document.add_paragraph())

        image = parser.parse_docx(
            docx_bytes(lambda d: (d.add_heading(HEADING, level=1), build(d)))
        ).images[0]
        assert image.data == PNG
        assert image.content_type == "image/png"
        assert image.name.endswith(".png")

    def test_a_picture_ahead_of_every_heading_is_dropped(self, docx_bytes):
        """It has no section to belong to, so it has no name either - and its bytes
        are never read, because naming happens only to blocks that were filed."""

        def build(document):
            _picture(document.add_paragraph())

        document = parser.parse_docx(docx_bytes(build))
        assert document.sections == []
        assert document.images == []

    def test_the_prefix_is_applied_where_the_document_is_rendered(self, docx_bytes):
        """The name in the content is bare, so the same parsed document can be
        rendered for S3, for a directory, or for anywhere else."""

        def build(document):
            _picture(document.add_paragraph())

        rendered = parser.parse_docx(
            docx_bytes(lambda d: (d.add_heading(HEADING, level=1), build(d)))
        ).markdown("https://example.org/docs/")
        assert "![](https://example.org/docs/1_img_1.png)" in rendered
