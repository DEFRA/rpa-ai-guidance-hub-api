"""What the conversion audit reads off the Word side of a document.

The audit is the oracle, so a mark it credits to the page that Word never draws is
not reported as a fault of its own: it is reported as a mark the parser lost, and
the search for it starts in the parser. These cases are about the Word side saying
what Word says and no more.

All fixture text is invented, as everywhere in the suite.
"""

from collections import Counter
from typing import Any

import audit_docx
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.text.paragraph import Paragraph


def _text_box(*runs: tuple[str, str]) -> Any:
    """A w:txbxContent holding one paragraph of runs, each with a colour or "" ."""
    box = OxmlElement("w:txbxContent")
    paragraph = OxmlElement("w:p")
    for text, colour in runs:
        run = OxmlElement("w:r")
        if colour:
            run.append(_colour(colour))
        element = OxmlElement("w:t")
        element.text = text
        run.append(element)
        paragraph.append(run)
    box.append(paragraph)
    return box


def _colour(value: str) -> Any:
    """The run properties that paint a run one colour."""
    properties = OxmlElement("w:rPr")
    element = OxmlElement("w:color")
    element.set(qn("w:val"), value)
    properties.append(element)
    return properties


def _in_style(document: Any, text: str, style_name: str) -> None:
    """Add a paragraph in a named style, defining the style if the template lacks it.

    python-docx will only apply a style the template already knows, and the styles
    worth testing against here - the contents styles a guide actually carries - are
    exactly the ones it does not have.
    """
    if all(style.name != style_name for style in document.styles):
        document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    document.add_paragraph(text, style=style_name)


def _bulleted(paragraph: Paragraph) -> None:
    """Put numbering on a paragraph, which is the whole of what makes it an item.

    The list itself is left undeclared: with no numbering part to read a format
    from, a bullet is what the audit reads, which is what Word draws when it has
    nothing else to draw.
    """
    numbering = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    numbering.get_or_add_numId().set(qn("w:val"), "1")


def _at_column(paragraph: Paragraph, left: int) -> None:
    """Draw a paragraph at a given indent, the way dragging one in Word does."""
    indent = paragraph._p.get_or_add_pPr().get_or_add_ind()
    indent.set(qn("w:left"), str(left))


def _item(document: Any, text: str, left: int) -> None:
    """A bulleted paragraph drawn at a given column, which is one list item."""
    paragraph = document.add_paragraph(text)
    _bulleted(paragraph)
    _at_column(paragraph, left)


def _anchor(paragraph: Paragraph, box: Any, colour: str = "") -> None:
    """Hang a text box off a run of the paragraph, the way a drawing does.

    Only the nesting matters to the walk, not the shape around it, so the w:drawing
    is written without the DrawingML that would tell Word how big to draw it.
    """
    run = paragraph.add_run()
    if colour:
        run._r.append(_colour(colour))
    drawing = OxmlElement("w:drawing")
    drawing.append(box)
    run._r.append(drawing)


class TestTextBoxMarks:
    def test_a_box_does_not_wear_the_marks_of_the_run_anchoring_it(self):
        """Word colours the anchor character, not the story the box holds.

        A real guide anchors a case note off a red run, and every word of the note
        is then read as red here while the page shows two of them that way. The
        parser marks up what the page shows, so the difference is charged to it: one
        section of that guide loses a seventh of its marks to a loss nobody made.
        """
        document = docx.Document()
        paragraph = document.add_paragraph()
        _anchor(
            paragraph,
            _text_box(("Case closed. ", ""), ("<input the date>", "FF0000")),
            colour="FF0000",
        )

        bag = audit_docx.Bag()
        audit_docx.mark_paragraph(bag, paragraph, frozenset())

        assert audit_docx.marks_of(bag.marks, audit_docx.RED) == Counter(
            {"input": 1, "the": 1, "date": 1}
        )

    def test_a_box_is_not_an_item_of_the_list_its_anchor_is_in(self):
        """A bulleted paragraph draws one bullet, and none on the box it anchors.

        A real guide bullets "Update the case note ... following the below template."
        and hangs the case note itself off that paragraph as a text box. Word draws
        no bullet on the box: its paragraphs carry no numbering of their own, and the
        same box drawn as a one-cell table - the same box, by this file's own
        reckoning of what a box is - is credited with no list at all. Inherited, the
        anchor's bullet charges the parser with seventy list marks that section never
        drew, in the one section of that guide scoring below 100%.
        """
        document = docx.Document()
        paragraph = document.add_paragraph("Update the case note.")
        _bulleted(paragraph)
        _anchor(paragraph, _text_box(("Example: Parcel ABC 1234.", "")))

        section = audit_docx.Section("Amending the agreement")
        audit_docx.absorb(section, paragraph, audit_docx.ListRun())

        assert audit_docx.marks_of(section.bag.marks, audit_docx.LIST) == Counter(
            {"update": 1, "the": 1, "case": 1, "note": 1}
        )
        assert audit_docx.marks_of(section.bag.marks, audit_docx.BOX) == Counter(
            {"example": 1, "parcel": 1, "abc": 1, "1234": 1}
        )


class TestContentsEntries:
    def test_a_stray_contents_entry_does_not_end_the_section(self):
        """A contents entry is navigation wherever it turns up, and no more than that.

        One guide carries an empty paragraph styled "Contents RPA" in the middle of a
        section. Ending the section there drops every paragraph after it up to the
        next heading - eighty-five words the page really shows, and every mark on
        them - so nothing is reported missing, because the Word side never claimed
        them, and the parser is charged instead with sixty-five marks it invented.
        An oracle that under-counts accuses; it does not report itself.

        Across all nineteen guides only this one paragraph carries a contents style
        after the first body heading. Every real contents entry sits ahead of it,
        where there is no open section to end, so closing one here was never the
        rule doing its job. `parser._body_paragraph` states the rule the other way
        round and is right: contents entries are left out wherever they turn up.
        """
        document = docx.Document()
        document.add_heading("Applying", level=1)
        document.add_paragraph("Before the entry.")
        _in_style(document, "", "Contents Entry")
        document.add_paragraph("After the entry.")

        sections = audit_docx.word_sections(document)

        assert [section.heading for section in sections] == ["Applying"]
        assert sections[0].bag.words == Counter(
            {"applying": 1, "before": 1, "the": 2, "entry": 2, "after": 1}
        )


class TestListSteps:
    """What one item did relative to the item before it.

    Depth itself is not scored, and cannot be: Word measures a level in twips and
    Markdown in columns, and both sides read depth relatively, so no absolute level
    on one side names the same thing as a level on the other. The step between two
    neighbouring items is the one statement each side can make in its own units, so
    it is the one the two are compared on.
    """

    def test_an_item_drawn_past_the_one_before_it_steps_in(self):
        document = docx.Document()
        document.add_heading("Working the case", level=1)
        _item(document, "If yes,", 720)
        _item(document, "Open the register", 1440)
        _item(document, "Read the status", 1440)

        [section] = audit_docx.word_sections(document)

        assert audit_docx.marks_of(section.bag.marks, audit_docx.LIST_INDENT) == (
            Counter({"open": 1, "the": 1, "register": 1})
        )
        assert not audit_docx.marks_of(section.bag.marks, audit_docx.LIST_OUTDENT)

    def test_an_item_drawn_back_towards_the_margin_steps_out(self):
        document = docx.Document()
        document.add_heading("Working the case", level=1)
        _item(document, "If yes,", 720)
        _item(document, "Open the register", 1440)
        _item(document, "If no,", 720)

        [section] = audit_docx.word_sections(document)

        assert audit_docx.marks_of(section.bag.marks, audit_docx.LIST_OUTDENT) == (
            Counter({"if": 1, "no": 1})
        )

    def test_an_item_dragged_a_little_is_still_in_its_neighbour_column(self):
        """Two lists pasted from different documents sit a few twips apart at the
        same depth. Read as a step, every such pair would report a nesting the page
        does not show, and the parser would be charged with flattening it."""
        document = docx.Document()
        document.add_heading("Working the case", level=1)
        _item(document, "If yes,", 714)
        _item(document, "If no,", 720)

        [section] = audit_docx.word_sections(document)

        assert not audit_docx.marks_of(section.bag.marks, audit_docx.LIST_INDENT)
        assert not audit_docx.marks_of(section.bag.marks, audit_docx.LIST_OUTDENT)

    def test_an_empty_paragraph_does_not_close_the_run(self):
        """Word spaces its lists with empty paragraphs and the parser writes no line
        for one, so a run closed here and not there puts the step on one side only -
        and the mark is then charged to whichever side was still counting."""
        document = docx.Document()
        document.add_heading("Working the case", level=1)
        _item(document, "If yes,", 720)
        document.add_paragraph("")
        _item(document, "Open the register", 1440)

        [section] = audit_docx.word_sections(document)

        assert audit_docx.marks_of(section.bag.marks, audit_docx.LIST_INDENT) == (
            Counter({"open": 1, "the": 1, "register": 1})
        )

    def test_the_markdown_side_reads_a_step_from_the_marker_column(self):
        bag = audit_docx.markdown_bag("- If yes,\n  - Open the register\n- If no,")

        assert audit_docx.marks_of(bag.marks, audit_docx.LIST_INDENT) == Counter(
            {"open": 1, "the": 1, "register": 1}
        )
        assert audit_docx.marks_of(bag.marks, audit_docx.LIST_OUTDENT) == Counter(
            {"if": 1, "no": 1}
        )

    def test_a_blank_line_does_not_close_the_markdown_run(self):
        """A blank line inside a list makes it loose rather than ending it, which is
        the same reason the Word side holds its run open across an empty paragraph."""
        bag = audit_docx.markdown_bag("- If yes,\n\n  - Open the register")

        assert audit_docx.marks_of(bag.marks, audit_docx.LIST_INDENT) == Counter(
            {"open": 1, "the": 1, "register": 1}
        )


class TestKnownLimits:
    """What the page shows that no Markdown could carry.

    Held out of the marks and counted apart, because the score exists to point at a
    difference somebody can go and fix. A loss the format makes unavoidable, left in
    the coverage column, reads as a fault nobody can repair and sits beside the ones
    that are real. `--missing` names each of these instead.
    """

    def test_a_list_in_a_table_cell_is_a_limit_rather_than_a_lost_mark(self):
        """A GFM pipe row cannot hold a newline, so `tables` joins a cell's blocks
        with <br> and its bullets become hyphens in the cell's text. No parser can
        do otherwise, and three of the guides would carry the shortfall for ever."""
        document = docx.Document()
        table = document.add_table(rows=1, cols=2)
        _bulleted(table.cell(0, 0).paragraphs[0])
        table.cell(0, 0).paragraphs[0].add_run("Open the register")

        section = audit_docx.Section("Working the case")
        audit_docx.absorb(
            section,
            table.cell(0, 0).paragraphs[0],
            audit_docx.ListRun(),
            frozenset({audit_docx.TABLE}),
        )

        assert not audit_docx.marks_of(section.bag.marks, audit_docx.LIST)
        assert audit_docx.marks_of(section.bag.limits, audit_docx.IN_A_CELL) == Counter(
            {"open": 1, "the": 1, "register": 1}
        )

    def test_a_list_in_a_callout_is_not_a_limit(self):
        """A box is a blockquote, and a blockquote holds blocks of its own - so a
        list inside one survives, and a shortfall there is the parser's to answer
        for. Only a pipe cell cannot carry it."""
        document = docx.Document()
        paragraph = document.add_paragraph("Open the register")
        _bulleted(paragraph)

        section = audit_docx.Section("Working the case")
        audit_docx.absorb(
            section, paragraph, audit_docx.ListRun(), frozenset({audit_docx.BOX})
        )

        assert audit_docx.marks_of(section.bag.marks, audit_docx.LIST) == Counter(
            {"open": 1, "the": 1, "register": 1}
        )
        assert not section.bag.limits

    def test_a_step_out_past_where_the_list_begins_is_a_limit(self):
        """A Markdown list has no column to the left of its first item, so a run
        opening indented and stepping back out cannot be drawn at all. The parser
        starts every run at the margin, which is the only thing it can do."""
        document = docx.Document()
        document.add_heading("Working the case", level=1)
        _item(document, "Open the register", 1440)
        _item(document, "Read the status", 1440)
        _item(document, "If no,", 720)

        [section] = audit_docx.word_sections(document)

        assert not audit_docx.marks_of(section.bag.marks, audit_docx.LIST_OUTDENT)
        assert audit_docx.marks_of(
            section.bag.limits, audit_docx.PAST_THE_START
        ) == Counter({"if": 1, "no": 1})

    def test_a_step_out_from_deeper_is_drawn_however_far_left_it_lands(self):
        """It is the item being left that decides. Leaving one drawn deeper than the
        run began, Markdown has an indent to bring back, and it draws the step even
        where the item arriving is further left than the run's own first item -
        because everything at or left of that is the margin, and the step from an
        indent to the margin is a step. Read the other way round, the audit calls a
        loss on eighty-eight marks the viewer plainly shows."""
        document = docx.Document()
        document.add_heading("Working the case", level=1)
        _item(document, "If yes,", 926)
        _item(document, "Open the register", 2203)
        _item(document, "If no,", 643)

        [section] = audit_docx.word_sections(document)

        assert audit_docx.marks_of(
            section.bag.marks, audit_docx.LIST_OUTDENT
        ) == Counter({"if": 1, "no": 1})
        assert not section.bag.limits

    def test_a_step_out_to_a_column_the_run_has_used_is_not_a_limit(self):
        """The run has been there, so Markdown has a depth to step back to and the
        parser is expected to draw it. This is the half that is still a fault."""
        document = docx.Document()
        document.add_heading("Working the case", level=1)
        _item(document, "If yes,", 720)
        _item(document, "Open the register", 1440)
        _item(document, "If no,", 720)

        [section] = audit_docx.word_sections(document)

        assert audit_docx.marks_of(
            section.bag.marks, audit_docx.LIST_OUTDENT
        ) == Counter({"if": 1, "no": 1})
        assert not section.bag.limits


class TestUncAddresses:
    r"""One address, two spellings, and the escape that separates them.

    Word writes `file:///\\host\share`; RFC 8089 writes `file://host/share`. They
    name one file, so the audit folds them together - otherwise a parser that spells
    an address correctly is scored as having lost it. What it must not fold is the
    address a renderer actually follows: `\\` between brackets is one backslash to
    CommonMark, so the characters written and the place they point are not the same
    thing, and reading them raw is how a broken link scores as a whole one.
    """

    def test_the_two_spellings_of_one_address_are_one_url(self):
        word = r"file:///\\server.example\share\Draft%20Letter"
        markdown = "file://server.example/share/Draft%20Letter"

        assert audit_docx.normalise_url(word) == audit_docx.normalise_url(markdown)

    def test_an_address_naming_another_file_is_another_url(self):
        word = r"file:///\\server.example\share\Draft%20Letter"

        assert audit_docx.normalise_url(word) != audit_docx.normalise_url(
            "file://server.example/share/Other%20Letter"
        )

    def test_the_canonical_spelling_matches_the_address_word_wrote(self):
        """What the parser now writes, scored against what Word holds."""
        word = r"file:///\\server.example\share\Draft%20Letter"
        bag = audit_docx.markdown_bag(
            "[Draft folder](file://server.example/share/Draft%20Letter)"
        )

        assert list(bag.urls) == [audit_docx.normalise_url(word)]

    def test_a_destination_is_read_with_its_escapes_spent(self):
        """The address is the one a reader is sent to, not the characters typed."""
        bag = audit_docx.markdown_bag(
            r"[Draft folder](file:///\\server.example\share\Draft%20Letter)"
        )

        assert list(bag.urls) == [r"file:///\server.example\share\draft%20letter"]

    def test_a_destination_a_renderer_mangles_is_not_the_address_word_wrote(self):
        """The loss the raw reading hid: one backslash short of the right server."""
        word = r"file:///\\server.example\share\Draft%20Letter"
        bag = audit_docx.markdown_bag(
            r"[Draft folder](file:///\\server.example\share\Draft%20Letter)"
        )

        assert list(bag.urls) != [audit_docx.normalise_url(word)]
