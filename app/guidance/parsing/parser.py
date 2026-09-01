"""Read a Word (.docx) package into a MarkdownDocument."""

from __future__ import annotations

import io
import re
import zipfile
from dataclasses import dataclass
from typing import TYPE_CHECKING, Any

import docx
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from app.guidance.parsing import images, inline, lists, models, tables, textboxes
from app.guidance.parsing.errors import DocumentParseError
from app.guidance.parsing.ooxml import is_toggle_on

if TYPE_CHECKING:
    from collections.abc import Iterator

    import docx.document

# Styles that open the document's navigation, so the cover page has ended.
_CONTENTS_STYLE_PREFIXES = ("toc", "contents", "table of contents")

# Word's own annex styles are not "Heading n", so they carry no level and have to be
# matched by name. Neither guide spells it "annex" or "schedule", but a template that
# did would mean exactly the same thing.
_APPENDIX_STYLE_PREFIXES = ("appendix", "annex", "schedule")

# An appendix is a top-level section. CS's contents pulls the style in with
# TOC \o "1-4" \h \z \t "Appendix,1" - the document declaring the level itself.
_APPENDIX_LEVEL = 1

# Word numbers its heading styles from 1, and the level is the whole of the name
# after the word: "Heading 2 Box" is a style in its own right, not a Heading 2.
_HEADING_STYLE = re.compile(r"heading\s*([1-9]\d*)$")

# The blocks of a document body. Anything else there - a bookmark, a section break,
# a proofing mark - says nothing the output carries.
_BODY_BLOCKS = (qn("w:p"), qn("w:tbl"))

_TITLE_STYLE = "title"

# Between the author's distinct title parts, and within one hand-wrapped part.
_TITLE_SEPARATOR = " — "
_WRAPPED_LINE_SEPARATOR = " "


def parse_docx(source: bytes) -> models.MarkdownDocument:
    """Parse the bytes of a .docx file into a MarkdownDocument.

    Raises:
        DocumentParseError: if the bytes are not a readable Word document.
    """
    document = _open(source)
    sections, bookmarks = _extract_sections(document)
    # Naming comes after the walk because only a finished section can say what its
    # pictures should be called, and a picture in a block that never reached a
    # section is dropped with the block rather than extracted and thrown away.
    images.name_all(sections, document.part)
    return models.MarkdownDocument(
        title=_extract_title(document),
        sections=sections,
        bookmarks=bookmarks,
    )


def _open(source: bytes) -> docx.document.Document:
    """Open .docx bytes, mapping every way that can fail to one error.

    python-docx raises a different exception for each kind of bad input, and
    only ever raises PackageNotFoundError for a *path* - never for the stream we
    hand it. The ValueError it raises for a non-Word package also interpolates the
    stream's repr into its message, which is meaningless to a caller, so the text
    is replaced here and the original kept as the cause.
    """
    try:
        return docx.Document(io.BytesIO(source))
    except zipfile.BadZipFile as exc:
        msg = "Source is not a .docx file: it is not a zip archive."
        raise DocumentParseError(msg) from exc
    except KeyError as exc:
        msg = "Source is not a .docx file: the zip archive is not an Office package."
        raise DocumentParseError(msg) from exc
    except (ValueError, PackageNotFoundError) as exc:
        msg = "Source is not a Word document."
        raise DocumentParseError(msg) from exc


def _extract_title(document: docx.document.Document) -> str:
    """Return the document's title, preferring the one printed on the page.

    The cover is read first and the document properties are only a fallback. Core
    properties are metadata that Word carries forward from whatever the file was
    copied from, so they go stale silently - a properties title can name a year the
    document was superseded from - while the cover is what a reader sees.
    """
    printed = _cover_title(document)
    if printed:
        return printed

    stored = document.core_properties.title
    return stored.strip() if stored else ""


def _cover_title(document: docx.document.Document) -> str:
    """Reconstruct the title printed on the cover page.

    Where the cover marks its title with the Title style, only those paragraphs are
    taken; otherwise the whole cover is. Each part is put back together as the author
    wrapped it, and the parts are joined with a dash. The text is joined but not
    otherwise tidied - a defect in a title is a finding to report, not noise to
    normalise away.
    """
    groups = _cover_groups(document)

    marked = [[p for p in group if _style_name(p) == _TITLE_STYLE] for group in groups]
    if any(marked):
        groups = [group for group in marked if group]

    return _TITLE_SEPARATOR.join(
        _WRAPPED_LINE_SEPARATOR.join(paragraph.text.strip() for paragraph in group)
        for group in groups
    )


def _cover_groups(document: docx.document.Document) -> list[list[Paragraph]]:
    """Return the cover page's paragraphs, grouped as the author laid them out.

    A blank paragraph separates one part of the title from the next; consecutive
    paragraphs are one part that the author wrapped by hand because it was too long
    for the line. Runs of blanks - including the ones padding the top of the page -
    separate but never form a part of their own. Headers and footers live in a
    separate XML part and so are excluded automatically.

    The cover ends at the first table of contents, body heading or page break. The
    two spellings of a page break stop it at different points, because they mean
    opposite things about the paragraph carrying them.
    """
    groups: list[list[Paragraph]] = []
    part: list[Paragraph] = []

    for paragraph in document.paragraphs:
        # This paragraph is already on the page after the cover, so its text is not
        # part of the title.
        if _opens_body(paragraph) or _starts_new_page(paragraph):
            break

        if paragraph.text.strip():
            part.append(paragraph)
        elif part:
            groups.append(part)
            part = []

        # ...whereas a break within the paragraph ends the page after it, so what it
        # says still belongs to the cover.
        if _ends_page(paragraph):
            break

    if part:
        groups.append(part)

    return groups


def _opens_body(paragraph: Paragraph) -> bool:
    """Whether this paragraph opens the navigation or the body, ending the cover.

    The heading check is what stops a document with no cover page break at all from
    swallowing its opening heading into the title. It asks the same question that
    opens a section, so the two rules cannot drift apart.
    """
    return (
        _is_contents(paragraph)
        or _is_appendix(paragraph)
        or _heading_level(paragraph) is not None
    )


def _is_contents(paragraph: Paragraph) -> bool:
    """Whether this paragraph is one of the document's contents entries."""
    return _style_name(paragraph).startswith(_CONTENTS_STYLE_PREFIXES)


def _starts_new_page(paragraph: Paragraph) -> bool:
    """Whether this paragraph is forced to the top of a new page by its properties."""
    properties = paragraph._p.find(qn("w:pPr"))
    if properties is None:
        return False
    return is_toggle_on(properties.find(qn("w:pageBreakBefore")))


def _ends_page(paragraph: Paragraph) -> bool:
    """Whether this paragraph carries an explicit page break in its own text."""
    return any(
        break_element.get(qn("w:type")) == "page"
        for run in paragraph.runs
        for break_element in run._r.findall(qn("w:br"))
    )


@dataclass
class _OpenSection:
    """A section still open for children as the walk moves down the document.

    `children` is how many sections have been opened directly beneath this one, and
    so is the ordinal the next one takes. The stack's first frame stands for the
    document itself, holding no section, so a top-level heading is counted and
    parented by the same code as any other.
    """

    level: int = 0
    section: models.MarkdownSection | None = None
    children: int = 0
    appendices: int = 0


def _extract_sections(
    document: docx.document.Document,
) -> tuple[list[models.MarkdownSection], dict[str, models.MarkdownSection]]:
    """Turn the document's headings into a flat list of sections in document order.

    Word does not put the number in the text. Numbering is attached to the heading
    *styles*, so Word generates "4.3.1.1" when it renders the page while the paragraph
    itself says only its title. The number is therefore derived here, from nothing but
    each heading's level relative to the one before it.

    Levels are read as relative, never absolute: a document that opens at Heading 2
    still starts at 1, and a heading that skips a level nests one deep rather than
    leaving a gap in the number. An appendix is the exception that proves it: its
    style declares a top level outright, and it is lettered rather than numbered.

    Every other paragraph is content, and belongs to the section opened most recently
    whatever its depth. Anything ahead of the first heading is not: what sits there
    is the cover page and the contents, and a contents page is regenerated from the
    headings anyway.

    Consecutive list items are the one kind of content that is not one paragraph to
    one block: a run of them is held open and rendered together, because what makes
    a Markdown list is the items standing next to each other. A table is the other:
    it is a block of the body in its own right, and closes any run open when it
    arrives.

    The bookmarks a cross-reference can point at are collected on the way past. One
    is claimed only where it marks the start of a section, that being the whole of
    what a number can be derived for. A bookmark landing anywhere else is better left
    as the raw name Word wrote than sent confidently to the wrong place.
    """
    sections: list[models.MarkdownSection] = []
    bookmarks: dict[str, models.MarkdownSection] = {}
    open_list: list[tuple[lists.ListItem, str]] = []
    # The document's own frame is never popped: it sits at level 0, and a heading's
    # level is never lower than 1.
    stack = [_OpenSection()]

    for element, opened_ahead_of_it in _body_items(document):
        if element.tag == qn("w:tbl"):
            _close_list(sections, open_list)
            _append_block(sections, tables.table_markdown(element, document))
            continue

        if element.tag == textboxes.TEXT_BOX:
            _close_list(sections, open_list)
            _append_block(sections, textboxes.markdown(element, document))
            continue

        paragraph = Paragraph(element, document)
        appendix = _is_appendix(paragraph)
        level = _APPENDIX_LEVEL if appendix else _heading_level(paragraph)
        heading = paragraph.text.strip()

        # A heading with nothing in it is a layout artefact - numbering it would put
        # a section in the output that the document does not have.
        if level is None or not heading:
            _collect_body(sections, open_list, paragraph)
            continue

        _close_list(sections, open_list)

        while stack[-1].level >= level:
            stack.pop()

        section = _open_beneath(stack[-1], heading, appendix=appendix)
        sections.append(section)
        stack.append(_OpenSection(level=level, section=section))

        for name in opened_ahead_of_it + _bookmark_names(paragraph._p):
            bookmarks[name] = section

    _close_list(sections, open_list)
    return sections, bookmarks


def _open_beneath(
    parent: _OpenSection, heading: str, *, appendix: bool
) -> models.MarkdownSection:
    """Open a section beneath `parent`, taking the next ordinal of its own kind.

    Appendices and numbered sections are counted apart, so an annex following
    section 7 is A rather than 8, and a numbered heading after that annex is 8.
    """
    if appendix:
        parent.appendices += 1
        ordinal = parent.appendices
    else:
        parent.children += 1
        ordinal = parent.children

    return models.MarkdownSection(
        heading=heading, ordinal=ordinal, parent=parent.section, appendix=appendix
    )


def _body_items(
    document: docx.document.Document,
) -> Iterator[tuple[Any, list[str]]]:
    """The body's blocks in document order, each with the bookmarks ahead of it.

    Elements are yielded as Word wrote them rather than as python-docx proxies, so
    that deciding what a block is stays with the caller and this walk has only the
    one job. `document.paragraphs` would return today's paragraphs and go on doing
    so, but it cannot see a table, and a table is a block of the body like any other.

    The walk also carries the bookmark names opened between one block and the next: a
    w:bookmarkStart marking a heading sits either inside that heading's paragraph or
    as a sibling just ahead of it, and a walk seeing only paragraphs would miss half
    of them. Names not claimed by the block that follows are dropped with it.
    """
    opened: list[str] = []
    for element in document.element.body:
        if element.tag == qn("w:bookmarkStart"):
            opened.extend(_bookmark_names(element))
        elif element.tag in _BODY_BLOCKS:
            yield element, opened
            opened = []
            if element.tag == qn("w:p"):
                # A text box is a block of the page that Word wrote inside a run
                # rather than beside it, so no walk of the body's own children can
                # reach it. Yielding it straight after the paragraph it hangs from is
                # what files it under the section that paragraph belongs to. It
                # claims no bookmarks: the ones opened here were the paragraph's.
                for box in textboxes.anchored_in(element):
                    yield box, []


def _bookmark_names(element: Any) -> list[str]:
    """Every bookmark name opened by this element or anything beneath it."""
    starts = element.iter(qn("w:bookmarkStart"))
    return [name for start in starts if (name := start.get(qn("w:name")))]


def _collect_body(
    sections: list[models.MarkdownSection],
    open_list: list[tuple[lists.ListItem, str]],
    paragraph: Paragraph,
) -> None:
    """Take one paragraph of the body: another item of the open list run, or prose.

    The list question is asked here rather than at the top of the walk, and that
    ordering is the whole of the guard on it: a document may attach numbering to its
    heading styles as well, so a rule reading numbering alone would bullet every
    heading in it. By the time a paragraph arrives here it is already not a heading.

    Contents entries are left out wherever they turn up, and are not list items
    however a document numbers them. They sit ahead of every heading in both real
    guides and so never reach here, but a document whose contents page opens with a
    heading of its own would otherwise file its whole table of contents as prose.
    """
    if _is_contents(paragraph):
        _close_list(sections, open_list)
        return

    item = lists.list_item(paragraph)
    if item is not None:
        open_list.append((item, inline.paragraph_markdown(paragraph)))
        return

    _close_list(sections, open_list)
    _append_block(sections, inline.paragraph_markdown(paragraph))


def _close_list(
    sections: list[models.MarkdownSection],
    open_list: list[tuple[lists.ListItem, str]],
) -> None:
    """File the open run of list items as one block, and open a fresh run.

    A run is closed by anything that is not a list item - a heading, a paragraph,
    or the end of the document - so it always lands in the section it started in.
    """
    if not open_list:
        return

    block = lists.render(open_list)
    open_list.clear()
    _append_block(sections, block)


def _append_block(sections: list[models.MarkdownSection], block: str) -> None:
    """Add one finished block of Markdown to the open section, where there is one.

    Anything ahead of the first heading has no section to belong to and is dropped.
    What sits there is the cover page and the contents, neither of which is content
    the conversion is meant to carry.
    """
    if not sections or not block:
        return

    section = sections[-1]
    section.content = f"{section.content}\n\n{block}" if section.content else block


def _is_appendix(paragraph: Paragraph) -> bool:
    """Whether this paragraph opens an appendix, which only its style can say.

    An annex style carries no outline level and is not a "Heading n", so there is
    nothing else to read it from. Detection is by style alone, deliberately: a
    "Heading n" whose text happens to read "Annex A" is already a section by the
    rule below, and would gain a letter rather than an existence.
    """
    return _style_name(paragraph).startswith(_APPENDIX_STYLE_PREFIXES)


def _heading_level(paragraph: Paragraph) -> int | None:
    """The level of a "Heading n" paragraph, or None if it is not one.

    The level has to come from the style name because that is the only place it is
    written: the outline level a document declares for a heading lives on the style
    definition, not on the paragraph. A custom style whose name only begins like a
    heading - "Heading Box", "Heading 2 Box" - names no level of its own and is
    therefore body text.
    """
    match = _HEADING_STYLE.match(_style_name(paragraph))
    return int(match[1]) if match else None


def _style_name(paragraph: Paragraph) -> str:
    """The paragraph's style name, lowercased, or "" when it has none.

    The two guards are not alike. python-docx annotates `style` as optional but
    resolves an unset or unknown style id to the document's default, so it never
    actually returns None - that guard is there for the type checker, and writing
    it as an expression keeps an unreachable branch out of the coverage report.
    `name` really can be None, for a style carrying no w:name element, which is
    covered by a test.
    """
    style = paragraph.style
    name = style.name if style is not None else None
    return name.lower() if name else ""
